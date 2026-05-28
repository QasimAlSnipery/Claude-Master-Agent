from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Romira R-ALA — 50 ChatGPT Image Ad Prompts', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    'Product: R Alpha Lipoic Acid Capsules 600mg by Romira | Target: Diabetic Neuropathy Sufferers\n'
    'Instructions: Copy each prompt into ChatGPT with the product image attached. '
    'ChatGPT will return an image generation prompt ready for DALL-E / Midjourney / Ideogram.'
)
doc.add_paragraph('')

PRODUCT_CONTEXT = (
    "I am attaching an image of our product: Romira R Alpha Lipoic Acid Capsules (600mg). "
    "This is a premium supplement that uses only the pure R-form of Alpha Lipoic Acid — the natural, body-identical isomer — at a clinical-strength 600mg dose. "
    "It is designed for people suffering from diabetic neuropathy: the burning, tingling, and numbness in the feet and hands caused by nerve damage from high blood sugar. "
    "R-ALA is the only antioxidant proven to cross the blood-brain barrier, regenerate vitamins C and E, improve insulin sensitivity, and directly relieve neuropathy symptoms. "
    "The product bottle is the centerpiece of the ad image. "
    "The brand is Romira — clean, premium, medical-trust aesthetic. "
    "Price: $39.99."
)

SUFFIX = (
    "The product image is attached — make sure the product bottle is prominently featured in the CENTER of the composition. "
    "Generate me a detailed image generation prompt I can paste directly into an AI image generator. "
    "The image MUST be 1:1 square ratio, ultra high quality, photorealistic, studio-grade lighting, sharp focus. "
    "Include the exact visual elements, lighting style, color palette, mood, and any text overlays in the prompt."
)

prompts = [
    # PAIN/SYMPTOM HOOKS
    (1, "Burning Feet at Night",
     f"{PRODUCT_CONTEXT} Create an ad for someone who wakes up at 2am with burning, tingling feet from diabetic neuropathy. "
     f"The mood is dark, nighttime bedroom scene — one person clutching their foot in pain on the left side, "
     f"and the Romira R-ALA bottle glowing with warm light on the right. Headline overlay: 'BURNING FEET AT NIGHT? YOUR NERVES NEED THIS.' {SUFFIX}"),

    (2, "The Fire Inside",
     f"{PRODUCT_CONTEXT} Create an ad visualizing nerve pain as literal fire or electricity running through the feet and legs, "
     f"with the Romira bottle standing calm and authoritative in the center. Color palette: deep red/orange fading to cool blue around the bottle. "
     f"Headline overlay: 'STOP THE FIRE FROM WITHIN.' {SUFFIX}"),

    (3, "Pins and Needles",
     f"{PRODUCT_CONTEXT} Create an ad showing a close-up of a diabetic person's foot covered in illustrated pins/needles (metaphor for neuropathy pain), "
     f"with the Romira bottle in the center foreground. Clean white background, clinical yet emotional. "
     f"Headline overlay: 'PINS & NEEDLES? THIS IS WHAT YOUR NERVES ARE MISSING.' {SUFFIX}"),

    (4, "Numb No More",
     f"{PRODUCT_CONTEXT} Create an ad split into two halves: left side shows a greyed-out, faded foot (numbness), right side shows the same foot vibrant and full of color. "
     f"Romira bottle centered between the two halves. Headline overlay: 'NUMB FEET AREN'T NORMAL. THEY'RE A WARNING.' {SUFFIX}"),

    (5, "Sleepless Pain",
     f"{PRODUCT_CONTEXT} Create a dark, moody nighttime ad. A digital clock reads 3:17am. A person's feet visible on white sheets, clearly uncomfortable. "
     f"The Romira bottle sits on the nightstand glowing softly. Headline: 'DIABETIC NEUROPATHY STEALS YOUR SLEEP. R-ALA GIVES IT BACK.' {SUFFIX}"),

    # PRODUCT SCIENCE ANGLES
    (6, "R vs S Form",
     f"{PRODUCT_CONTEXT} Create an educational-style ad comparing R-ALA vs S-ALA. Left side: a generic grey supplement bottle labeled 'S-FORM — SYNTHETIC'. "
     f"Right side: Romira bottle labeled 'R-FORM — PURE & NATURAL'. Clean white background, medical aesthetic. "
     f"Headline: 'NOT ALL ALA IS THE SAME. ONLY R-FORM WORKS.' {SUFFIX}"),

    (7, "The Mitochondria Ad",
     f"{PRODUCT_CONTEXT} Create a futuristic, scientific-looking ad showing illustrated mitochondria glowing with energy in the background, "
     f"with the Romira bottle centered in the foreground. Dark background, electric blue and gold energy particles. "
     f"Headline: '600MG OF THE MOLECULE YOUR MITOCHONDRIA HAVE BEEN WAITING FOR.' {SUFFIX}"),

    (8, "Blood-Brain Barrier",
     f"{PRODUCT_CONTEXT} Create an ad showing a stylized human brain with a glowing shield/barrier around it, "
     f"and R-ALA molecules (illustrated as small golden dots) passing through. Romira bottle centered at the bottom. "
     f"Headline: 'THE ONLY ANTIOXIDANT THAT REACHES YOUR BRAIN.' {SUFFIX}"),

    (9, "Antioxidant Recycler",
     f"{PRODUCT_CONTEXT} Create an ad showing circular arrows (recycling symbol) made of illustrated vitamins C and E, "
     f"with the Romira bottle at the center. Clean, bright, scientific. "
     f"Headline: 'R-ALA RECYCLES VITAMINS C & E — DOUBLING YOUR PROTECTION.' {SUFFIX}"),

    (10, "Clinical Dose",
     f"{PRODUCT_CONTEXT} Create a clean, clinical white ad showing the Romira bottle center-stage with '600MG' in large bold type above it. "
     f"A subtle graph line in the background showing nerve function improvement over weeks. "
     f"Headline: 'THE DOSE ACTUALLY USED IN CLINICAL STUDIES.' {SUFFIX}"),

    # EMOTIONAL/LIFESTYLE
    (11, "Walk Without Pain",
     f"{PRODUCT_CONTEXT} Create an uplifting ad showing a senior walking comfortably on a beach or park path, feet clearly visible and healthy, "
     f"with the Romira bottle in the foreground bottom-center. Warm golden hour lighting. "
     f"Headline: 'WALK WITHOUT THINKING ABOUT YOUR FEET AGAIN.' {SUFFIX}"),

    (12, "Back to Life",
     f"{PRODUCT_CONTEXT} Create an emotional before/after ad. Before: person sitting alone, looking down at their feet in pain, grey tones. "
     f"After: same person laughing with family, active, vibrant colors. Romira bottle centered between the two scenes. "
     f"Headline: 'NEUROPATHY TOOK THEIR LIFE. R-ALA GAVE IT BACK.' {SUFFIX}"),

    (13, "The First Step",
     f"{PRODUCT_CONTEXT} Create a powerful ad showing a close-up of a bare foot taking a confident first step on clean white ground, "
     f"with the Romira bottle prominently displayed. Tagline: 'THE FIRST STEP TO PAIN-FREE NERVES STARTS HERE.' {SUFFIX}"),

    (14, "Grandparent Active",
     f"{PRODUCT_CONTEXT} Create a warm, emotional ad showing a grandparent playing with their grandchildren — running, laughing — feet and legs clearly healthy. "
     f"Romira bottle in foreground. Soft warm lighting. Headline: 'DON'T LET NEUROPATHY TAKE YOU OUT OF THE MOMENTS THAT MATTER.' {SUFFIX}"),

    (15, "Independence Ad",
     f"{PRODUCT_CONTEXT} Create an empowering ad showing a diabetic senior standing tall and confident, arms slightly open, radiating independence. "
     f"Romira bottle centered below. Headline: 'DIABETIC NEUROPATHY DOESN'T HAVE TO WIN.' {SUFFIX}"),

    # URGENCY/FEAR
    (16, "Progression Warning",
     f"{PRODUCT_CONTEXT} Create an urgent medical-style ad with a timeline showing neuropathy progression: tingling → burning → numbness → permanent damage. "
     f"A red arrow stops at 'PERMANENT DAMAGE' with an X through it. Romira bottle centered below. "
     f"Headline: 'NEUROPATHY GETS WORSE EVERY DAY YOU WAIT.' {SUFFIX}"),

    (17, "67% Stat Ad",
     f"{PRODUCT_CONTEXT} Create a bold stat-based ad. Large number '67%' dominates the top half. "
     f"Romira bottle centered in the bottom half. Headline: '67% OF DIABETICS DEVELOP NEUROPATHY. MOST DON'T ACT UNTIL IT'S TOO LATE.' {SUFFIX}"),

    (18, "Nerve Death Clock",
     f"{PRODUCT_CONTEXT} Create an urgent ad showing a stylized nerve cell with a fading/dying effect — edges going dark. "
     f"Romira bottle centered, glowing with health. Headline: 'EVERY MONTH WITHOUT R-ALA, YOUR NERVES PAY THE PRICE.' {SUFFIX}"),

    (19, "Before It's Permanent",
     f"{PRODUCT_CONTEXT} Create a stark, high-contrast black and white ad except the Romira bottle which is in full color. "
     f"Headline in bold red: 'NERVE DAMAGE CAN BECOME PERMANENT. ACT NOW.' {SUFFIX}"),

    (20, "The Window Is Closing",
     f"{PRODUCT_CONTEXT} Create an ad showing an hourglass with sand running out, and inside the glass the imagery transitions from healthy nerves to damaged ones. "
     f"Romira bottle centered at the base of the hourglass. Headline: 'THERE'S STILL TIME TO PROTECT YOUR NERVES.' {SUFFIX}"),

    # SOCIAL PROOF / TESTIMONIAL STYLE
    (21, "5-Star Review",
     f"{PRODUCT_CONTEXT} Create a dark premium ad showing a 5-star rating in gold at the top, a short quote: "
     f"'The burning in my feet is almost completely gone after 3 weeks.' — Verified buyer, Type 2 diabetic. "
     f"Romira bottle centered below. Clean black background, gold accents. {SUFFIX}"),

    (22, "Community Testimonial",
     f"{PRODUCT_CONTEXT} Create a warm social-proof ad showing 3 small profile photo placeholders (diverse, 50s-70s age range) each with a short quote about nerve relief, "
     f"Romira bottle prominently centered. Headline: 'THOUSANDS OF DIABETICS ALREADY SLEEP BETTER.' {SUFFIX}"),

    (23, "Before After Testimonial",
     f"{PRODUCT_CONTEXT} Create a split testimonial ad: left side quote — 'Week 1: Still burning every night.' Right side — 'Week 4: I forgot I had neuropathy.' "
     f"Romira bottle centered. Star rating underneath the Week 4 quote. {SUFFIX}"),

    (24, "Doctor Endorsed Style",
     f"{PRODUCT_CONTEXT} Create a clinical trust ad. A person in a white coat (no face shown, just coat and hands) holding the Romira bottle. "
     f"Clean white/light blue background. Headline: 'R-ALA IS THE MOST STUDIED COMPOUND FOR DIABETIC NERVE DAMAGE.' {SUFFIX}"),

    (25, "Real Results Timeline",
     f"{PRODUCT_CONTEXT} Create an ad showing a simple 4-week timeline with icons: Week 1 — less burning, Week 2 — better sleep, Week 3 — more sensation, Week 4 — significant relief. "
     f"Romira bottle centered below the timeline. Headline: 'WHAT 4 WEEKS OF R-ALA LOOKS LIKE.' {SUFFIX}"),

    # COMPARISON/EDUCATION
    (26, "Why Cheap ALA Fails",
     f"{PRODUCT_CONTEXT} Create an educational comparison ad. Left side: generic bottle labeled 'RACEMIC ALA (50% useless S-form)' with a red X. "
     f"Right side: Romira bottle labeled 'PURE R-FORM — 100% BIOACTIVE' with a green checkmark. "
     f"Headline: 'YOU'VE BEEN BUYING THE WRONG ALA.' {SUFFIX}"),

    (27, "Insulin Sensitivity",
     f"{PRODUCT_CONTEXT} Create a clean, clinical ad showing a glucose meter and the Romira bottle side by side. "
     f"Illustrated arrows showing blood sugar dropping to healthy range. "
     f"Headline: 'R-ALA IMPROVES INSULIN SENSITIVITY — AND PROTECTS YOUR NERVES AT THE SAME TIME.' {SUFFIX}"),

    (28, "Fat AND Water Soluble",
     f"{PRODUCT_CONTEXT} Create a scientific-visual ad showing the human body silhouette with both fat cells and water cells illuminated, "
     f"R-ALA molecules (golden dots) reaching both. Romira bottle centered at bottom. "
     f"Headline: 'THE ONLY ANTIOXIDANT THAT WORKS IN EVERY CELL OF YOUR BODY.' {SUFFIX}"),

    (29, "Root Cause Ad",
     f"{PRODUCT_CONTEXT} Create an ad with a split concept: top half shows 'MASKING SYMPTOMS' (pain pills, creams, socks) crossed out, "
     f"bottom half shows 'FIXING THE ROOT CAUSE' with Romira bottle centered and glowing. "
     f"Headline: 'STOP MANAGING. START HEALING.' {SUFFIX}"),

    (30, "The Molecule Behind the Relief",
     f"{PRODUCT_CONTEXT} Create a dark, dramatic ad showing the chemical structure of R-Alpha Lipoic Acid illustrated beautifully in gold lines, "
     f"Romira bottle centered and glowing in the foreground. "
     f"Headline: 'THIS MOLECULE HAS BEEN STUDIED FOR 30 YEARS. NOW IT'S IN ONE CAPSULE.' {SUFFIX}"),

    # SPECIFIC SYMPTOM ADS
    (31, "Tingling Hands",
     f"{PRODUCT_CONTEXT} Create an ad focused on hand neuropathy — showing illustrated electric/tingling sensation around the hands. "
     f"Romira bottle centered, hands reaching toward it. "
     f"Headline: 'TINGLING HANDS? IT'S NOT JUST CIRCULATION. YOUR NERVES NEED R-ALA.' {SUFFIX}"),

    (32, "Numbness Ad",
     f"{PRODUCT_CONTEXT} Create an ad showing a foot/hand transitioning from grey/faded (numb) to full color and sensation, "
     f"Romira bottle at center. Headline: 'NUMBNESS IS YOUR NERVES DYING. GIVE THEM WHAT THEY NEED.' {SUFFIX}"),

    (33, "Electric Shock Pain",
     f"{PRODUCT_CONTEXT} Create an ad showing illustrated electric shock bolts (yellow/white) around a leg/foot, "
     f"then the Romira bottle standing calm and powerful in the center, the electricity fading near it. "
     f"Headline: 'ELECTRIC SHOOTING PAINS? YOUR NERVES ARE INFLAMED.' {SUFFIX}"),

    (34, "Cold Feet",
     f"{PRODUCT_CONTEXT} Create an ad showing blue/ice-cold illustrated feet on the left, warm and normal feet on the right, "
     f"Romira bottle centered. Headline: 'ALWAYS COLD FEET? THAT'S NERVE DAMAGE TALKING.' {SUFFIX}"),

    (35, "Balance Problems",
     f"{PRODUCT_CONTEXT} Create an ad showing a person carefully walking, slightly unstable, neuropathy affecting balance. "
     f"Romira bottle in foreground bottom-center, grounding the composition. "
     f"Headline: 'NEUROPATHY ISN'T JUST PAIN. IT STEALS YOUR BALANCE TOO.' {SUFFIX}"),

    # PREMIUM / BRAND TRUST
    (36, "Clean Label",
     f"{PRODUCT_CONTEXT} Create a minimalist premium ad. Pure white background. Romira bottle dead center, beautifully lit with studio lighting. "
     f"Three clean benefit lines below: ✓ 600mg Pure R-Form | ✓ No Fillers | ✓ Clinical-Strength. "
     f"Headline above: 'THE PUREST R-ALA ON THE MARKET.' {SUFFIX}"),

    (37, "Premium Studio Shot",
     f"{PRODUCT_CONTEXT} Create a pure product-hero ad. Romira bottle on a dark marble surface, dramatic side lighting, "
     f"soft reflections, moody premium feel. Single headline above: 'CLINICAL-STRENGTH NERVE SUPPORT.' {SUFFIX}"),

    (38, "Gold Standard",
     f"{PRODUCT_CONTEXT} Create a luxury ad with deep black background, gold particle effects surrounding the Romira bottle. "
     f"Bottle lit with warm golden light. Headline in gold serif font: 'THE GOLD STANDARD IN NEUROPATHY SUPPORT.' {SUFFIX}"),

    (39, "Made Different",
     f"{PRODUCT_CONTEXT} Create a confidence ad with 'NOT ALL SUPPLEMENTS ARE EQUAL' at the top. "
     f"Generic grey bottle on the left fading into shadow. Romira bottle on the right, fully illuminated. "
     f"Headline below: 'ROMIRA. MADE FOR RESULTS.' {SUFFIX}"),

    (40, "Trust Badges",
     f"{PRODUCT_CONTEXT} Create a clean trust-building ad. Romira bottle centered with badges arranged around it: "
     f"'600mg Clinical Dose' | 'Pure R-Form Only' | 'No Synthetic Fillers' | 'GMP Certified'. "
     f"White background, premium feel. Headline: 'EVERYTHING YOUR NERVES NEED. NOTHING THEY DON'T.' {SUFFIX}"),

    # VALUE / OFFER ADS
    (41, "Price Value Ad",
     f"{PRODUCT_CONTEXT} Create a value-framing ad. Romira bottle centered. "
     f"Comparison text: 'Pain clinic visit: $300+ | Romira R-ALA: $39.99/month'. "
     f"Headline: 'REAL NERVE SUPPORT DOESN'T HAVE TO COST A FORTUNE.' {SUFFIX}"),

    (42, "30-Day Challenge",
     f"{PRODUCT_CONTEXT} Create an energetic challenge-style ad. Bold headline: '30-DAY NEUROPATHY CHALLENGE'. "
     f"Romira bottle centered with a calendar/timer element. "
     f"Subheadline: 'Take R-ALA for 30 days. Feel the difference or your money back.' {SUFFIX}"),

    (43, "Subscribe & Save",
     f"{PRODUCT_CONTEXT} Create a clean value ad showing the Romira bottle with a subscription badge '— SAVE 20% —'. "
     f"Headline: 'CONSISTENT RELIEF. CONSISTENT SAVINGS.' Simple, premium, white background. {SUFFIX}"),

    (44, "Risk-Free Ad",
     f"{PRODUCT_CONTEXT} Create a reassurance-focused ad. Romira bottle centered with a shield icon overlay saying 'MONEY-BACK GUARANTEE'. "
     f"Headline: 'TRY IT RISK-FREE. YOUR NERVES HAVE NOTHING TO LOSE.' {SUFFIX}"),

    (45, "Fast Relief Ad",
     f"{PRODUCT_CONTEXT} Create an urgency-meets-hope ad. Romira bottle centered with 'RELIEF AS FAST AS 2 WEEKS' "
     f"in bold type above. Clean white background with subtle green health-color accents. {SUFFIX}"),

    # CREATIVE / CONCEPTUAL
    (46, "Nerve Network",
     f"{PRODUCT_CONTEXT} Create a stunning visual ad showing a glowing blue nerve network map (like a city seen from above at night) "
     f"with the Romira bottle as the power source at the center, energy radiating outward. "
     f"Headline: 'POWER YOUR NERVE NETWORK.' {SUFFIX}"),

    (47, "The Shield",
     f"{PRODUCT_CONTEXT} Create a powerful ad showing a glowing shield made of light protecting an illustrated nerve cell, "
     f"with the Romira bottle at the bottom powering the shield. Dark background, electric blue and gold. "
     f"Headline: 'PROTECT YOUR NERVES AT THE CELLULAR LEVEL.' {SUFFIX}"),

    (48, "Sugar vs Nerves",
     f"{PRODUCT_CONTEXT} Create a dramatic visual of sugar crystals/cubes attacking nerve cells, "
     f"with the Romira bottle acting as a barrier between them. Headline: 'HIGH BLOOD SUGAR DESTROYS NERVES. R-ALA FIGHTS BACK.' {SUFFIX}"),

    (49, "Mitochondria Power Plant",
     f"{PRODUCT_CONTEXT} Create a cinematic ad showing a glowing mitochondria (illustrated beautifully like a power plant) "
     f"with energy beams flowing outward and the Romira bottle centered in the foreground. "
     f"Headline: 'FUEL YOUR CELLS. FREE YOUR NERVES.' {SUFFIX}"),

    (50, "The Simple Truth",
     f"{PRODUCT_CONTEXT} Create a bold, stripped-down typographic ad. Clean white background. Romira bottle centered. "
     f"Three lines of text only: 'YOUR NERVES ARE DAMAGED.' / 'THEY CAN HEAL.' / 'START WITH R-ALA.' "
     f"No clutter. Maximum impact. {SUFFIX}"),
]

for num, name, prompt_text in prompts:
    doc.add_heading(f'Prompt {num}: {name}', level=2)
    p = doc.add_paragraph(prompt_text)
    p.style.font.size = Pt(11)
    doc.add_paragraph('')

output_path = r'C:\Users\Darin Game\Downloads\Romira_RALA_50_ChatGPT_Ad_Prompts.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
