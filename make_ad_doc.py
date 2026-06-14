# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# base style
st = doc.styles['Normal']
st.font.name = 'Calibri'
st.font.size = Pt(11)

def title(t):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

def h(t):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x88, 0x6a, 0x00)
    p.space_before = Pt(10)

def body(t):
    for line in t.split('\n'):
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(2)

def numlist(items):
    for i, it in enumerate(items, 1):
        doc.add_paragraph(f"{i}. {it}")

# ============ DOC TITLE ============
title("Prestige — Meta Ad Copy")
doc.add_paragraph("Primary text + headlines for the Jacob Daisy & Naughty O'Clock watches.").italic = True
doc.add_paragraph("")

# ============ JACOB DAISY ============
title("1) Jacob Daisy Dragon Tourbillon")

h("Primary Text")
body(
"Most watches try to stand out…\n"
"This one doesn't try — it commands the room.\n"
"\n"
"A double-axis flying tourbillon spins on the dial. A rose-gold dragon circles a turquoise Earth. "
"The hours are read through a daisy — and the caseback says it all: “The World Is Mine.”\n"
"\n"
"47mm of pure presence. 172 stones of shine.\n"
"The kind of watch people stop you in the street to ask about.\n"
"\n"
"Limited stock. Don't miss it."
)

h("Recommended Headline")
body("The World Is Mine.")

h("Headline Options (10)")
numlist([
 "The World Is Mine.",
 "A Universe On Your Wrist",
 "600+ Parts. Two Spinning Axes. One Wrist.",
 "The Most Complex Watch You'll Ever Own",
 "Watch The Tourbillon Come Alive",
 "Not A Watch — A Statement",
 "Built For Men Who Refuse Average",
 "The Daisy. The Dragon. The Tourbillon.",
 "Engineered To Be Watched",
 "Limited Edition. Unlimited Attention.",
])

doc.add_page_break()

# ============ NAUGHTY OCLOCK ============
title("2) Naughty O'Clock Iced Rainbow Watch")

h("Primary Text")
body(
"Some watches tell time.\n"
"This one tells everyone exactly what kind of person you are.\n"
"\n"
"Iced-out rainbow bezel. Heavy steel build. And a dial so cheeky people do a double-take "
"every single time they look.\n"
"\n"
"The ultimate conversation starter — and the gift that gets the biggest reaction at every party.\n"
"\n"
"4 bold colors. Limited stock. Grab yours before they're gone."
)

h("Recommended Headline")
body("The Watch Everyone Asks About")

h("Headline Options (10)")
numlist([
 "The Watch Everyone Asks About",
 "Naughty O'Clock — Tell Me You Didn't Look Twice",
 "Bold. Iced. Unapologetic.",
 "The Cheekiest Watch On The Internet",
 "Look Closer… We Dare You",
 "The Best Gag Gift That's Secretly A Real Flex",
 "Rainbow Ice. Naughty Dial. Zero Regrets.",
 "The Conversation Starter On Your Wrist",
 "Some Watches Are Boring. This One Isn't.",
 "Iced-Out & A Little Bit Naughty",
])

doc.save("Prestige_Ad_Copy.docx")
print("saved Prestige_Ad_Copy.docx")
