from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading('Sandra Hughes — Final 4 Facebook Posts', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

note = doc.add_paragraph('These 4 posts feed into: romira-pea-advertorial-updated.html\nAdvertorial headline: "Your Body Already Makes Its Own Pain-Off Switch... But After 50, It Stops Making Enough"')
note.alignment = WD_ALIGN_PARAGRAPH.CENTER

posts = [
    {
        'number': '1 — LEAD POST',
        'topic': 'Grandkids (Run This First)',
        'why': 'LEAD POST: Deepest emotional hook. Grandkids = the #1 motivational stake for 50+ women. Highest share and comment rate. Matches "pain-off switch" angle perfectly — she found the reason nothing worked.',
        'headline': 'I sat on the floor with my granddaughter. I actually cried.',
        'image': '''Ultra-photorealistic, ultra-high-quality candid photograph of a British woman aged 55-58 seated in a worn faded floral armchair in a modest English living room. She is leaning slightly forward, both hands resting on her knees, looking down at the floor in front of her with a soft, bittersweet expression — the look of someone who wants to do something but cannot. In the foreground on the carpet, slightly out of her reach, a half-finished wooden jigsaw puzzle with bright colourful pieces. A child's pink cardigan is visible draped over the arm of the sofa nearby, suggesting a grandchild was just here. The woman is wearing a soft sage green V-neck knit jumper and dark navy trousers. Her hair is shoulder-length, a natural mix of warm brown and silver-grey, slightly loose. The room shows real domestic life — family photographs clustered on the wall behind her in mismatched frames, a bookshelf with paperbacks, a small side table with a half-drunk mug of tea and a TV remote. Natural late-afternoon window light from the left creates warm but slightly dim shadows. Slightly off-centre composition. Photorealistic DSLR quality with shallow depth of field. No studio lighting. No filters. Looks like her son-in-law took it quietly from the doorway. Aspect ratio 1:1 square.''',
        'primary': '''Evie is 4. She wanted me down on the floor doing her puzzle with her this afternoon.

I had to tell her Nana's back was being poorly again.

She just nodded. Like she already knew. And that absolutely broke me.

I have tried everything. Physio, glucosamine, turmeric, the copper bracelet my mum swore by. My GP told me it was just age and to stay active. I wanted to laugh in his face.

Then I came across an article that actually explained WHY nothing was working. Not a sales pitch. A proper explanation of something my body used to make naturally — that apparently slows right down after 50. I had never heard of it before.

I gave it five weeks. Last Tuesday, Evie asked me to sit with her and do the puzzle. And I got down on that floor, did the whole puzzle, and got back up.

I went to the kitchen and cried. A happy cry.

If you have been dismissed by your doctor and told to just live with it — read this before you give up. Link in the comments. 👇🙏'''
    },
    {
        'number': '2',
        'topic': 'Morning Back Pain (Widest Reach)',
        'why': 'WIDEST AUDIENCE: Every 50+ person with pain knows the morning struggle. Maximum cold traffic reach. The "pain-off switch" angle answers exactly why mornings are the worst.',
        'headline': 'My back finally stopped ruining my mornings',
        'image': '''Ultra-photorealistic, ultra-high-quality candid UGC-style photograph of a British woman aged 54-58 sitting on the absolute edge of an unmade bed in a modest English bedroom, early morning. She is in the act of bracing herself to stand — one hand pressed firmly into the small of her back, the other hand gripping the mattress edge, weight shifted forward. Her expression is a tight-lipped wince — brow deeply furrowed, eyes slightly squinted downward, mouth closed and compressed. She is wearing a faded dusty-pink cotton nightshirt, hair dishevelled from sleep, slightly pushed back on one side. The bedding is mismatched floral — slightly tangled, one pillow fallen sideways. On the wooden nightstand: a rectangular seven-day pill organiser (the rectangular kind from the chemist), a glass of water with a small straw, and a crumpled tissue. Warm yellow bedside lamp is on. Through the net curtain, grey early English morning light creates cool ambient light from one side, warm lamp from the other. Slight motion blur on her hand. Horizon very slightly tilted. The image looks exactly like she propped her phone against the lamp and triggered it herself. No studio lighting. No ring-light eye reflection. Photorealistic DSLR quality. Aspect ratio 1:1 square.''',
        'primary': '''Four years of dreading mornings.

Not just discomfort. I mean genuinely lying there before I even tried to move, steeling myself for it. That first attempt to stand used to take me 10 minutes.

Physio, heat pads, ibuprofen, the foam mattress my daughter made me buy. Helpful for a day or two. Nothing lasted.

I came across an article that explained something I had never heard before — that your body actually produces its own natural pain relief, and that after 50 it starts making less and less of it. That is why mornings get worse as the years go on. Not just age. An actual biological reason.

I started taking something to support it. I did not expect much honestly.

By week four I was downstairs making tea before my husband woke up. No grabbing the bedpost. No 10-minute ritual.

It has been three months. I have not had a morning like the old ones since.

The article that explained it all is linked below — worth a read if mornings are hard for you too. 👇'''
    },
    {
        'number': '3',
        'topic': 'Painsomnia / 3AM (Highest Intent)',
        'why': 'HIGHEST BUYER INTENT: People awake at 3am scrolling are at peak desperation. The "pain-off switch stops working at night" angle matches the advertorial mechanism perfectly.',
        'headline': 'I slept through the night for the first time in 2 years',
        'image': '''Ultra-photorealistic, ultra-high-quality intimate bedroom photograph taken at 3am. A British woman aged 55-60 is sitting fully upright in bed, back against the headboard, knees slightly drawn up. She is wearing an oversized faded white cotton sleep shirt, no makeup, hair flattened from the pillow on one side. The only light source is a warm amber-yellow bedside lamp creating uneven, intimate shadows across the room. Her expression shows complete exhausted resignation — eyes half-open and heavy-lidded, brow slightly furrowed, mouth loosely closed and slightly downturned. One hand rests on her knee. On the nightstand: a digital alarm clock displaying 3:14 in red numerals, a glass of water with condensation on the outside, a small amber pill bottle with the label facing away, and a paperback book lying face-down. The bedsheets are twisted and pushed to one side — evidence of tossing and turning. Drugstore reading glasses sit folded on the nightstand. The room is slightly cluttered — a cardigan draped over the chair in the corner, a tissue box. Photorealistic DSLR quality with visible film grain from low light. Slightly blurry as if taken on an older phone in the dark. Aspect ratio 1:1 square.''',
        'primary': '''Who else is awake right now?

It is 3am and the burning in my feet and the aching in my lower back woke me up again. Same as every night for the past two years.

I am exhausted all day and then I cannot sleep. My doctor said try to rest more. I actually laughed.

About six weeks ago I read an article that changed how I understood what was happening to me at night. Apparently the natural compound your body makes to calm nerve and joint pain — it dips dramatically at night after 50. That is why pain gets so much worse when you lie still. It is not in your head. There is a biological reason the nights are harder.

I started supporting my body with it. Week one nothing dramatic. Week three I woke up and it was morning. Actual morning. I had slept through.

I sat there for a moment not quite believing it.

If you are part of the 3am club tonight — please read this before you write it off as just getting older. The link is in the comments. 👇🙏❤️'''
    },
    {
        'number': '4',
        'topic': 'Tried Everything (Most Ready to Buy)',
        'why': 'MOST READY TO BUY: People who have spent money trying to fix their pain and failed are the most motivated buyers. The "reason nothing worked" angle of the advertorial is the perfect answer to this post.',
        'headline': 'After £600 and 14 things that failed — this finally made sense',
        'image': '''Ultra-photorealistic, ultra-high-quality overhead flat-lay photograph of a real British woman's cluttered wooden bedside table, taken directly from above on a phone. On the surface, arranged as if genuinely lived-in rather than styled: a seven-day pill organiser with some compartments open, a tube of Deep Heat cream with the cap left off, a copper magnetic bracelet with slight tarnish, a half-empty bottle of turmeric capsules (Holland and Barrett style label), an NHS physiotherapy exercise sheet printed on A4 paper with some exercises circled in blue biro, an ice pack in a faded blue fabric cover, and a half-finished mug of chamomile tea with the teabag still in it. One weathered hand is reaching into the frame from the lower right — visible age spots on the back of the hand, short unpainted nails, a simple gold wedding band — and is picking up the turmeric bottle. The table surface has a fine layer of dust in the corners. Warm amber light from a bedside lamp creates soft uneven shadows. The whole scene reads exactly as: someone who has been trying everything for a long time. Slight phone lens barrel distortion at the edges. No artificial styling. Photorealistic DSLR quality. Aspect ratio 1:1 square.''',
        'primary': '''I kept a list.

Turmeric. Glucosamine. Deep Heat. The copper bracelet my mum swore by. Physio exercises every morning for three months. Ice. Heat. Ibuprofen. Cutting out dairy. Cutting out gluten. CBD oil. Epsom salt baths. Stretches from YouTube.

Fourteen things. Two years. Roughly £600.

Some helped a tiny bit for a few days. Nothing lasted. Nothing got to the root of it.

I was ready to accept this was just my life now.

Then I found an article — not a sales pitch, an actual explanation — that told me WHY nothing I had tried was working. It was not that the things were useless. It was that none of them addressed what was actually happening in my body after 50. There is a compound your body produces naturally to regulate pain and inflammation. After 50, production drops significantly. Everything I had been taking was treating the surface. Not the cause.

It took me six weeks to admit it was working. Because I did not want to get my hopes up again.

But it is working. It is the only thing that has come off my list not because it stopped — but because it actually worked.

Article is linked in the comments. Read it before you spend another penny on anything else. 👇'''
    }
]

for post in posts:
    doc.add_page_break()

    doc.add_heading(f"POST {post['number']} — {post['topic']}", level=1)

    why_para = doc.add_paragraph()
    why_run = why_para.add_run(post['why'])
    why_run.bold = True
    why_run.font.color.rgb = RGBColor(0x2C, 0x6E, 0x8A)

    doc.add_paragraph()

    doc.add_heading('HEADLINE', level=2)
    p = doc.add_paragraph(post['headline'])
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(14)

    doc.add_paragraph()

    doc.add_heading('IMAGE PROMPT (paste into Midjourney / DALL-E / Gemini)', level=2)
    doc.add_paragraph(post['image'])

    doc.add_paragraph()

    doc.add_heading('PRIMARY TEXT (Facebook Post Caption)', level=2)
    doc.add_paragraph(post['primary'])

    doc.add_paragraph()
    doc.add_paragraph('--- DROP YOUR ADVERTORIAL LINK IN THE FIRST COMMENT ---')
    doc.add_paragraph()

output_path = r'C:\Users\moham\OneDrive\Desktop\Sandra_Final4_Posts.docx'
doc.save(output_path)
print('Done:', output_path)
