from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading('Romira PEA Advertorial — Image Prompts', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Copy and paste each prompt directly into Midjourney, DALL-E, or Firefly. Do not change anything.')
doc.add_paragraph('')

prompts = [
    (
        'IMAGE 1 — Hero (Top of Article) | 16:9 — EMOTIONAL PAIN SCENE',
        'A photorealistic candid photograph of a woman aged 55-60 years old, sitting on the edge of her bed in the early morning, hunched slightly forward, both hands pressing into her knees with visible effort as she tries to push herself up to stand. She has light-to-medium skin with natural deep age lines, silver-streaked dark brown hair dishevelled from sleep, no makeup, eyes showing exhaustion and quiet frustration — not screaming in pain, but the expression of someone who wakes up every single morning dreading this exact moment. She is wearing a pale blue or cream cotton pyjama set, fabric creased and real. Her posture is the story — rounded shoulders, weight shifted forward, jaw slightly tightened, one foot flat on the floor and the other barely touching — the body language of someone summoning effort just to stand up. The bedroom around her is real and lived-in: unmade bed behind her with rumpled white linen, soft morning light coming through half-open curtains creating a warm but heavy atmosphere, a glass of water and a small orange pill bottle on the nightstand beside her, slightly blurred in the background. Shot from a low angle slightly to the side — like a documentary photographer crouched in the corner of the room. The frame is wide enough to show her full upper body and the bed behind her. Shallow depth of field — her face and hands are razor sharp, background softly blurred. Color palette: muted blues, pale creams, warm morning amber from the window. 35mm lens feel, Canon 5D documentary quality, natural film grain, no studio lighting whatsoever, no stock photo cleanliness. Ultra-realistic skin — visible pores, dark under-eyes, natural sleep creases on her face. No AI skin smoothing. This image must feel like you walked into her bedroom uninvited and caught a private moment she would never show anyone. Aspect ratio 16:9. --ar 16:9 --v 6 --style raw'
    ),
    (
        'IMAGE 2 — Failed Solutions Flat Lay | 4:3',
        'A photorealistic close-up flat lay photograph of a cluttered bathroom or kitchen countertop showing failed supplement and medication attempts. Items scattered naturally — not arranged for a photoshoot — include: an open orange prescription-style ibuprofen bottle with a few white tablets spilled out beside it, a small dark amber glass CBD oil dropper bottle with the label partially peeled at one corner, a half-used jar of turmeric powder with a light yellow dusting of powder around its lid, and an open blister pack of glucosamine tablets with two missing from the pack. The counter surface is white marble with grey veining, slightly dusty and used. Natural overhead light from above, slightly cool daylight tone. One or two pills have rolled slightly apart from the bottle. The scene feels like someone has tried everything and left it all out in defeat. Shot from directly above (birds-eye flat lay), extremely sharp focus across all items, no depth of field blur. Slight real-world imperfections — fingerprint smudge on the CBD bottle, powder residue on the counter. No branding visible or legible on any label. Color palette: white counter, amber glass, orange plastic, muted yellows, clinical whites. Ultra-photorealistic, no illustration, no cartoon, no clean studio composition. Aspect ratio 4:3. --ar 4:3 --v 6 --style raw'
    ),
    (
        'IMAGE 3 — Relief / Mechanism Moment | 4:3',
        'A photorealistic candid photograph of a woman aged 52-58 years old standing up from a cushioned armchair in a bright living room, caught mid-motion in a natural moment of ease. She has medium skin tone, dark brown shoulder-length hair with natural grey highlights, wearing a soft cream-coloured linen blouse tucked into straight dark navy trousers. Her posture is upright and relaxed — not straining, not grimacing — she looks like the movement is effortless and unremarkable to her, perhaps reaching for a book on a side table or just standing to walk away. Her expression is neutral and calm, slightly distracted, not performing. The living room around her is warm and tasteful — a cream fabric armchair, soft beige rug, warm afternoon light coming through large windows to the right side, casting long warm golden shadows across the wooden floor. Shot from a medium distance at eye level, like a candid lifestyle photographer in the corner of the room. Background slightly blurred but still recognizable. Warm tones — cream, oak, soft amber light. 50mm lens feel, natural grain, no studio lighting, no posed stock photo energy. Ultra-realistic skin, natural body proportions for a 55-year-old woman. Aspect ratio 4:3. --ar 4:3 --v 6 --style raw'
    ),
    (
        'IMAGE 4 — Product Shot (Romira PEA Bottle) | 1:1',
        'A professional ultra-clean product photography shot of a single white supplement bottle centered on a pure white background. The bottle is cylindrical, approximately 4 inches tall, with a white body and a deep navy blue (#264165) label. The label reads "Romira PEA 600mg" in clean modern sans-serif typography. The cap is white and slightly domed. Lighting is soft studio three-point lighting — a key light from the upper left, a fill light from the right, and a subtle backlight creating a gentle rim glow around the bottle edges. The bottle casts a very soft, barely-visible shadow directly beneath it on the white surface. The surface beneath the bottle is a seamless pure white sweep background — no texture, no gradients. The bottle label is sharply in focus from top to bottom. Shot straight-on at eye level, perfectly centered. No hands, no props, no lifestyle elements. Clinical, premium, pharmaceutical-adjacent feel — not trendy, not colorful, not playful. Similar aesthetic to high-end supplement brands. Ultra-sharp product photography, 85mm macro lens quality. Aspect ratio 1:1. --ar 1:1 --v 6 --style raw'
    ),
    (
        'IMAGE 5 — Sarah Story Image | 1:1',
        'A photorealistic candid UGC-style photograph of a woman aged 55-60 years old, walking down three or four outdoor stone or concrete steps — like front porch steps of a suburban home — in the late morning. She has light skin, warm-toned silver and brown hair cut to the chin, wearing a relaxed dark olive-green zip-up fleece jacket and straight mid-wash blue jeans, white sneakers. She is mid-step, weight on her leading foot, one hand loosely touching the metal handrail but not gripping it — like she is touching it out of habit, not necessity. Her expression is calm, slightly distracted, with the faintest trace of a relaxed smile — like someone who forgot for a moment that stairs used to be a problem. Shot from slightly below and to the side, handheld feel, as if a friend took the photo on a smartphone. Slight motion micro-blur on the trailing foot to convey real movement. Natural daylight, slight overcast — soft even light, no harsh shadows. Background: blurred suburban front yard, green hedges, a parked car partially visible. Color palette: warm olive, mid-blue denim, cool silver hair, muted green garden. Ultra-realistic, no retouching, no studio, genuine skin texture for a 57-year-old woman. Aspect ratio 1:1. --ar 1:1 --v 6 --style raw'
    ),
    (
        'IMAGE 6A — Testimonial Avatar: Margaret, Age 63 | 1:1',
        "A photorealistic close-up portrait headshot of a woman aged 62-65 years old. She has fair skin with natural age lines across her forehead, around her eyes (crow's feet), and smile lines around her mouth — all visible and unreduced. Her hair is fully silver-white, cut in a soft layered bob to the jaw. She is wearing a simple dusty rose or muted burgundy top, just the collar visible at the bottom of the frame. Her expression is warm and genuine — a soft, slightly tired smile, like someone sharing an honest review, not posing. Eyes are light blue or green, slightly crinkled at the corners when she smiles. Shot from chest-up, centered, facing slightly three-quarters toward the camera. Background is a soft blurred warm beige or cream interior — completely out of focus, no detail recognizable. Lighting is soft natural window light from the left. No makeup or minimal natural makeup. Real skin — pores, texture, natural imperfections visible. Ultra-photorealistic, no AI skin smoothing, no illustration. Aspect ratio 1:1. --ar 1:1 --v 6 --style raw"
    ),
    (
        'IMAGE 6B — Testimonial Avatar: Claire, Age 51 | 1:1',
        'A photorealistic close-up portrait headshot of a woman aged 49-53 years old. She has medium-warm skin tone, dark chestnut brown hair with a few natural grey strands at the temples, hair falling straight to the shoulders. She is wearing a simple white or light blue fitted top, just the collar visible. Her expression is bright and direct — a genuine confident smile, like someone who has good news and wants to share it. Eyes are dark brown, sharp and clear, with natural under-eye definition. Slight natural laugh lines around the mouth. Shot from chest-up, centered, facing almost straight to camera with a very slight head tilt. Background is soft blurred light grey or soft white — completely out of focus. Lighting is clean natural daylight from slightly above and to the right. Minimal makeup — mascara, light lip color only. Real skin texture, natural pores, no AI smoothing. Ultra-photorealistic. Aspect ratio 1:1. --ar 1:1 --v 6 --style raw'
    ),
    (
        'IMAGE 6C — Testimonial Avatar: David, Age 58 | 1:1',
        'A photorealistic close-up portrait headshot of a man aged 56-60 years old. He has medium-fair skin with natural age lines across his forehead, around his eyes, and along his jaw. Short salt-and-pepper hair — mostly grey at the temples, darker brown on top — neatly but not perfectly groomed. Light stubble across the jaw and upper lip, 2-3 days of growth, slightly grey. He is wearing a dark navy or charcoal grey collared shirt, collar just visible. His expression is calm and understated — a quiet, closed-mouth half-smile, like a man giving an honest straightforward review. Eyes are blue-grey, steady and direct. Shot from chest-up, centered, facing three-quarters toward the camera. Background is soft blurred warm grey or muted blue-grey interior — completely out of focus. Lighting is soft even natural light, slightly diffused. Real skin — visible pores, natural lines, no retouching. Ultra-photorealistic, no AI smoothing, no illustration. Aspect ratio 1:1. --ar 1:1 --v 6 --style raw'
    ),
]

for label, prompt in prompts:
    h = doc.add_heading(label, level=1)
    h.runs[0].font.color.rgb = RGBColor(0x26, 0x41, 0x65)
    p = doc.add_paragraph(prompt)
    p.runs[0].font.size = Pt(11)
    p.runs[0].font.name = 'Courier New'
    doc.add_paragraph('')

output = 'C:/Users/Darin Game/Downloads/Romira_PEA_Image_Prompts.docx'
doc.save(output)
print('Saved:', output)
