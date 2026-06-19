from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

def heading(text, size=16, color=(180, 60, 0), after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(*color)
    return p

def label(p, text, bold=True):
    r = p.add_run(text)
    r.bold = bold
    return r

# ---- title ----
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("ROMIRA ALPHA LIPOIC ACID")
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(180, 60, 0)

st = doc.add_paragraph()
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st.add_run('Claymation Ad Storyboard - "The Nerves Who Needed a Hero"')
r.italic = True; r.font.size = Pt(13)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Total length: ~60 seconds  |  6 clips x 10 seconds  |  Style: stop-motion claymation  |  Audience: adults 55+  |  For AI video generation")
r.font.size = Pt(10); r.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

heading("HOW TO USE THIS DOCUMENT", 13, (60, 60, 60))
doc.add_paragraph(
    "Each clip is 10 seconds long. Every clip has a NARRATION SCRIPT of about 18-20 words - this is the "
    "exact wording the voiceover should read in those 10 seconds. The pace is slow, warm, and clear "
    "because the audience is adults aged 55 and older. The first 5 clips educate; the final clip sells. "
    "The 'What is happening' breakdown describes the visuals so you can hand each clip to ChatGPT and ask "
    "it to write a detailed AI-video prompt (for Kling, Runway, Sora, or Pika). Keep the claymation look "
    "consistent across every clip."
)

heading("VISUAL STYLE - KEEP THIS THE SAME IN EVERY CLIP", 13, (60, 60, 60))
for b in [
    "Medium: handmade modeling-clay / plasticine, stop-motion claymation.",
    "Nerve characters: stubby clay creatures with long tail-like axons, big expressive eyes. Healthy = soft gold glow. Inflamed = swollen, angry red-orange, sparking at the tips.",
    "Villains (oxidative stress / free radicals): tiny jagged grey-black spiky clay sparks that buzz and poke.",
    "Hero (Alpha Lipoic Acid): a smooth round clay molecule with a warm golden glow, friendly face, small arms.",
    "Lighting: warm, cinematic, soft shadows. Camera: smooth slow moves, shallow depth of field.",
]:
    p = doc.add_paragraph(style="List Bullet"); p.add_run(b)

doc.add_paragraph()

clips = [
    {
        "title": "CLIP 1 - ZOOM IN: FROM OUTSIDE THE LEG TO INSIDE (0:00-0:10)",
        "summary": "Cinematic zoom from the outside of a human leg, through the skin, down to the inflamed nerves deep inside.",
        "beats": [
            "0-2s: Wide shot of a person's leg (clay style) standing in a soft warm-lit clay world. A subtle red glow pulses under the skin, hinting at trouble.",
            "2-4s: Camera pushes in fast toward the lower leg / calf. Clay skin texture fills the frame.",
            "4-6s: Camera passes THROUGH the skin surface - a smooth dive inward, layers of clay tissue peeling past.",
            "6-8s: We arrive inside the leg: a soft squishy clay tunnel. First glimpse of the nerve bundles running through it.",
            "8-10s: Settle on the nerves - they are swollen, glowing angry red-orange, clearly inflamed and twitching.",
        ],
        "narration": "Deep inside your legs are tiny nerves. As we age, they can swell, ache, and slowly lose their protection.",
        "words": 19,
    },
    {
        "title": "CLIP 2 - THE INFLAMED NERVES IN PAIN (0:10-0:20)",
        "summary": "Close-up on the suffering nerve characters. They are inflamed, sparking, and clearly hurting.",
        "beats": [
            "10-12s: Close-up on 3-4 nerve creatures. They are swollen and red, tips flashing small red sparks.",
            "12-14s: Nerve 1 winces and clutches itself. A red shockwave ripples down its axon tail.",
            "14-16s: Nerve 2 shivers - a visual cue of tingling and numbness (small static-like clay flecks).",
            "16-18s: The whole bundle pulses red in unison, like a heartbeat of pain.",
            "18-20s: Nerve 1 looks toward camera, eyes worried, and speaks.",
        ],
        "narration": "When nerves get inflamed, you feel it - burning, tingling, numbness, that pins-and-needles ache in your feet and hands.",
        "words": 20,
    },
    {
        "title": "CLIP 3 - THE VILLAINS: OXIDATIVE STRESS (0:20-0:30)",
        "summary": "The cause is revealed - swarms of tiny grey spiky free-radical sparks attacking the nerves.",
        "beats": [
            "20-22s: Pull back slightly. Tiny jagged grey-black spiky sparks (free radicals) buzz into frame.",
            "22-24s: The sparks poke and jab the nerves. Each jab makes a nerve flare brighter red.",
            "24-26s: More and more villains swarm in - the screen fills with chaotic grey buzzing.",
            "26-28s: The nerves shrink back, huddling together, dimming, overwhelmed.",
            "28-30s: Nerve 3 shouts out in fear.",
        ],
        "narration": "The cause? Tiny invaders called free radicals. Every single day, they attack your nerves and wear down their natural defenses.",
        "words": 20,
    },
    {
        "title": "CLIP 4 - THE CRY FOR SUPPORT (0:30-0:40)",
        "summary": "The nerves are at their lowest point and call out for help. The mood is dark right before the turn.",
        "beats": [
            "30-32s: The nerves are huddled, dim, drooping. Lighting goes cooler and darker.",
            "32-34s: One nerve sinks down, almost giving up. Soft, sad tone.",
            "34-36s: Nerve 1 lifts its head and looks up, searching for something.",
            "36-38s: A tiny glimmer of gold appears far off in the tunnel - hope on the way.",
            "38-40s: The nerves notice the light; their eyes widen.",
        ],
        "narration": "Left unprotected, your nerves grow weaker over time. They need real support - a strong shield against this daily, ongoing damage.",
        "words": 20,
    },
    {
        "title": "CLIP 5 - THE HERO ARRIVES: ALPHA LIPOIC ACID (0:40-0:50)",
        "summary": "The glowing golden ALA molecule rolls in, scatters the villains, and protects the nerves.",
        "beats": [
            "40-42s: The golden ALA molecule rolls in, radiating warm light. The tunnel brightens.",
            "42-44s: ALA sweeps through - its glow scatters and dissolves the grey free-radical sparks.",
            "44-46s: ALA turns to the nerves with a friendly smile and a little wave.",
            "46-48s: A warm golden shield-glow spreads over each nerve, wrapping them in protection.",
            "48-50s: ALA speaks, confident and kind.",
        ],
        "narration": "Meet Alpha Lipoic Acid - a powerful antioxidant. It fights those free radicals and helps protect your nerves, naturally.",
        "words": 18,
    },
    {
        "title": "CLIP 6 - RELIEF + PRODUCT + CTA (0:50-1:00)",
        "summary": "The nerves heal and glow gold again, then the clay Romira bottle appears with the call to action.",
        "beats": [
            "50-52s: The red inflammation fades. The nerves return to a calm, healthy gold glow.",
            "52-54s: Nerves stretch tall, stop sparking, and smile - the tingling is gone.",
            "54-56s: Camera pulls back as the whole nerve tunnel glows warm and healthy.",
            "56-58s: A clay-style Romira Alpha Lipoic Acid bottle rises into frame; the nerve characters cheer around it.",
            "58-60s: On-screen text appears: 'Romira Alpha Lipoic Acid - Support your nerves, naturally.' with 'romira.store'.",
        ],
        "narration": "Give your nerves the support they deserve. Try Romira Alpha Lipoic Acid today - visit romira.store and feel the difference.",
        "words": 20,
    },
]

for c in clips:
    heading(c["title"], 14, (180, 60, 0), after=2)
    p = doc.add_paragraph(); label(p, "Summary: "); p.add_run(c["summary"])

    p = doc.add_paragraph()
    label(p, "NARRATION (read slowly - " + str(c["words"]) + " words):")
    p = doc.add_paragraph()
    r = p.add_run('"' + c["narration"] + '"')
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0, 90, 40)

    p = doc.add_paragraph(); label(p, "What is happening on screen (second-by-second):")
    for b in c["beats"]:
        bp = doc.add_paragraph(style="List Bullet"); bp.add_run(b)
    doc.add_paragraph()

heading("TIP FOR CHATGPT PROMPTING", 13, (60, 60, 60))
doc.add_paragraph(
    "Paste one clip at a time into ChatGPT and say: 'Write a detailed claymation AI-video prompt for "
    "this 10-second clip, including camera movement, lighting, character look, action, and mood. "
    "Keep the stop-motion clay style consistent.' Generate all 6, then stitch the clips in order."
)

out = os.path.join(os.path.expanduser("~"), "Downloads", "Romira_ALA_Claymation_Storyboard_v2.docx")
doc.save(out)
print("SAVED:", out)
