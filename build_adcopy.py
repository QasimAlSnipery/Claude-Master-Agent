from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

def heading(text, size=15, color=(180, 60, 0)):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(size); r.font.color.rgb = RGBColor(*color)
    return p

# title
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("ROMIRA ALPHA LIPOIC ACID - AD COPY")
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor(180, 60, 0)
st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st.add_run("Primary Text + Headlines for the Claymation Nerve-Health Video Ad  |  Audience: adults 55+")
r.italic = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(100, 100, 100)
doc.add_paragraph()

# ---------- HEADLINES ----------
heading("10 HEADLINE OPTIONS (the short bold line under the video)")
headlines = [
    "Support Your Nerves, Naturally",
    "Tingling Feet? Your Nerves Need This",
    "Calm the Tingling. Support Your Nerves.",
    "Give Your Nerves the Support They Deserve",
    "The Tiny Molecule Your Nerves Have Been Missing",
    "Burning, Tingling, Numb Feet? Start Here",
    "Why Your Nerves Feel Tired After 55",
    "Protect Your Nerves From Daily Damage",
    "One Antioxidant. Healthier, Happier Nerves.",
    "Your Nerves Work Hard. Help Them Back.",
]
for i, h in enumerate(headlines, 1):
    p = doc.add_paragraph()
    r = p.add_run(str(i) + ". "); r.bold = True
    r = p.add_run(h); r.bold = True; r.font.size = Pt(12)
doc.add_paragraph()

# ---------- PRIMARY TEXTS ----------
heading("10 PRIMARY TEXT OPTIONS (the caption above the video)")

primary = [
    # 1 - story / educational
    "Deep inside your legs run tiny, delicate nerves. As the years pass, they can swell, tingle, and slowly lose their natural protection - leaving you with that burning, numb, pins-and-needles feeling.\n\n"
    "The cause? Free radicals quietly wearing your nerves down, day after day.\n\n"
    "Romira Alpha Lipoic Acid is a powerful antioxidant that helps fight free radicals and support healthy nerves - naturally.\n\n"
    "Give your nerves the support they deserve. -> romira.store",

    # 2 - direct problem-agitate-solve
    "If your feet burn, tingle, or go numb... your nerves are trying to tell you something.\n\n"
    "As we age, oxidative stress wears down the protection around our nerves - and the discomfort slowly creeps in.\n\n"
    "Alpha Lipoic Acid is one of nature's most powerful antioxidants for nerve support. Romira makes it simple to give your nerves the daily help they need.\n\n"
    "Feel the difference. -> romira.store",

    # 3 - empathy / relatable
    "Trouble sleeping because of tingling legs? Wincing on your walks? You are not alone - and it's not just 'getting older.'\n\n"
    "Behind that discomfort are tired, unprotected nerves. They need real support.\n\n"
    "Romira Alpha Lipoic Acid helps fight the daily damage and supports healthy nerve function, so you can get back to the simple days you love.\n\n"
    "-> romira.store",

    # 4 - question hook
    "Did you know your nerves are under attack every single day?\n\n"
    "Tiny invaders called free radicals slowly wear down the natural defenses that keep your nerves healthy - and that's when the burning, tingling, and numbness begin.\n\n"
    "Romira Alpha Lipoic Acid is a powerful antioxidant that helps shield your nerves and support how they feel and function.\n\n"
    "Protect what keeps you moving. -> romira.store",

    # 5 - short and punchy
    "Burning. Tingling. Numb.\n\n"
    "Your nerves need backup - and Alpha Lipoic Acid is one of the most trusted antioxidants for nerve support.\n\n"
    "Romira makes it easy. Support your nerves, naturally. -> romira.store",

    # 6 - benefit-led
    "Healthy nerves mean better sleep, easier walks, and more comfortable days.\n\n"
    "Romira Alpha Lipoic Acid delivers a powerful antioxidant that fights free radicals and supports your nerves from within.\n\n"
    "Thousands of people over 55 are giving their nerves the support they deserve. Join them. -> romira.store",

    # 7 - educational authority
    "Alpha Lipoic Acid is a unique antioxidant - it works in both water and fat, helping protect nerve cells throughout the body.\n\n"
    "As we age, our natural levels drop, and our nerves lose some of their protection. That's where Romira comes in.\n\n"
    "One simple step a day to support healthy nerves. -> romira.store",

    # 8 - reassurance / gentle
    "It's never too late to take care of your nerves.\n\n"
    "That tingling and numbness doesn't have to be 'just part of aging.' With the right antioxidant support, you can help protect the nerves that keep you active and independent.\n\n"
    "Romira Alpha Lipoic Acid - gentle, natural, daily nerve support. -> romira.store",

    # 9 - story open loop
    "Watch what's really happening inside your legs...\n\n"
    "Your nerves are swelling, sparking, and slowly losing their protection - and free radicals are the reason why.\n\n"
    "See how one powerful antioxidant, Alpha Lipoic Acid, helps fight back and support your nerves.\n\n"
    "Romira. Support your nerves, naturally. -> romira.store",

    # 10 - urgency / care
    "Every day you wait, free radicals keep wearing your nerves down.\n\n"
    "The good news? You can start supporting them today. Romira Alpha Lipoic Acid is a powerful antioxidant that helps fight that daily damage and supports healthy nerve function.\n\n"
    "Give your nerves the support they deserve - starting now. -> romira.store",
]

for i, p_text in enumerate(primary, 1):
    hp = doc.add_paragraph()
    r = hp.add_run("OPTION " + str(i)); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0, 90, 40)
    for block in p_text.split("\n\n"):
        doc.add_paragraph(block)
    doc.add_paragraph()

# ---------- RECOMMENDATIONS ----------
heading("MY TOP RECOMMENDATIONS")
recs = [
    ("Primary Text Option 1 + Headline 4 (\"Give Your Nerves the Support They Deserve\")",
     "Best match to the video. The primary text mirrors the storyboard's narration, so the ad and video feel like one piece - strong for cold 55+ traffic that needs education before the sell."),
    ("Primary Text Option 3 + Headline 2 (\"Tingling Feet? Your Nerves Need This\")",
     "Most emotional and relatable. Leads with real-life pain (sleep, walking) which 55+ buyers feel instantly. High scroll-stop and click intent."),
    ("Primary Text Option 5 + Headline 6 (\"Burning, Tingling, Numb Feet? Start Here\")",
     "Short, punchy, symptom-first. Great as a second variation to test against the longer educational version - good for retargeting warm audiences."),
]
for title, why in recs:
    p = doc.add_paragraph(); r = p.add_run(title); r.bold = True; r.font.size = Pt(12)
    doc.add_paragraph(why)
    doc.add_paragraph()

p = doc.add_paragraph()
r = p.add_run("Note: keep wording to nerve SUPPORT and COMFORT (not disease cure claims) to stay compliant with Meta ad policy for supplements.")
r.italic = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(120, 120, 120)

out = os.path.join(os.path.expanduser("~"), "Downloads", "Romira_ALA_Ad_Copy.docx")
doc.save(out)
print("SAVED:", out)
