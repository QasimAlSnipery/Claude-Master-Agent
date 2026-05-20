from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('R-Alpha Lipoic Acid — Facebook Ad Copy', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')

# ── HEADLINES ──────────────────────────────────────────────────────────────────
doc.add_heading('10 HEADLINES', level=1)

headlines = [
    "If Your Feet Burn, Tingle, or Go Numb at Night — This Is Why",
    "The Antioxidant That Works Where Vitamin C Can't Reach",
    "Most ALA Supplements Use the Wrong Form. Here's the Difference.",
    "Why Do Your Feet Burn at Night? (And What Actually Helps)",
    "Your Nerves Need More Than Just Magnesium",
    "600mg. Pure R-Form. Coconut Oil for Absorption. Finally, an ALA That Does the Job.",
    "The \"Universal Antioxidant\" — Works Inside AND Outside Your Cells",
    "Tingling. Burning. Pins and Needles. Your Nerves Are Telling You Something.",
    "R vs S: The Form of ALA You Take Changes Everything",
    "If You've Tried Everything for Nerve Discomfort — You Probably Missed This",
]

for i, h in enumerate(headlines, 1):
    p = doc.add_paragraph()
    p.add_run(f"{i}. {h}").bold = True

rec = doc.add_paragraph()
rec.add_run("⭐ TOP PICKS: Headlines #1, #10").bold = True
rec.runs[0].font.color.rgb = RGBColor(0x26, 0x41, 0x65)

doc.add_paragraph('')

# ── PRIMARY TEXTS ───────────────────────────────────────────────────────────────
doc.add_heading('10 PRIMARY TEXTS', level=1)

ads = [
    (
        "1. The Night Story",
        """That burning, tingling feeling in your feet. Worst at night. When the room is quiet and there's nothing to distract you from it.

You've probably tried things. Magnesium. CBD. Turmeric. Maybe nothing really worked the way you hoped.

Here's what most people don't know: nerves are some of the most energy-hungry, oxidative-stress-sensitive cells in your body. They need a very specific kind of support — not a general antioxidant, but one that works inside AND outside the cell membrane.

That's R-Alpha Lipoic Acid.

It's both water- and fat-soluble — which means it protects nerve cells in places Vitamin C and Vitamin E simply can't reach. It also supports mitochondrial energy and healthy blood sugar — two things directly connected to long-term nerve health.

Romira uses the pure R-form (the form your body actually produces) paired with Organic Coconut Oil so it absorbs properly.

600mg. No fillers. 60-day guarantee.

→ Try it risk-free"""
    ),
    (
        "2. The Form Matters",
        """If you've ever taken an alpha-lipoic acid supplement and felt nothing — there's a good reason.

Most ALA products use a 50/50 mix of R and S forms. The S-form is synthetic. It's not found in nature. Your body doesn't produce it.

The R-form is what your body actually makes. It's the biologically active enantiomer — the one that works.

Romira R-Alpha Lipoic Acid uses pure R-form ALA at 600mg, paired with Organic Coconut Oil because ALA is fat-soluble and needs it to absorb properly.

If your feet tingle. If your hands burn. If nerve discomfort is interfering with sleep, walking, or just getting through the day —

Give your nerves the right form.

60-day money-back guarantee. Made in the USA."""
    ),
    (
        "3. The Universal Antioxidant",
        """Vitamin C is water-soluble. It protects the watery parts of your cells.

Vitamin E is fat-soluble. It protects the fatty membranes.

R-Alpha Lipoic Acid is both.

That's why researchers call it the "universal antioxidant." It's the only antioxidant that can do its job both inside cells and around them — which is exactly why it matters so much for nerve tissue.

Nerves are surrounded by fatty myelin sheaths. They operate in energy-dense, oxidative environments. Most antioxidants can only protect part of the picture.

R-ALA protects all of it. And it regenerates Vitamin C, Vitamin E, and Glutathione while it's at it.

Romira R-Alpha Lipoic Acid. 600mg. Pure R-form. Organic Coconut Oil for absorption.

→ 60-day guarantee. Try it."""
    ),
    (
        "4. Short & Punchy (Mobile-First)",
        """Burning feet at night. Tingling hands. Pins and needles that won't quit.

That's your nerves under oxidative stress.

R-Alpha Lipoic Acid is the only antioxidant that works inside AND outside your cells — protecting nerve tissue where others can't reach.

Most brands use a cheap 50/50 mix of R and S forms. The S-form is synthetic. Romira uses pure R-form ALA + Organic Coconut Oil for absorption.

600mg. No fillers. 60-day money-back guarantee.

Your nerves have been waiting for this."""
    ),
    (
        "5. The Energy Angle",
        """Your nerve cells are the most energy-demanding cells in your body.

They need a constant supply of mitochondrial energy to send signals, maintain function, and stay resilient under stress. When that energy drops — when oxidative stress builds up — nerves become irritated, sensitive, and uncomfortable.

R-Alpha Lipoic Acid supports mitochondrial energy production at the cellular level. It also fights oxidative stress across water AND fat-soluble environments — something no other single antioxidant can do.

That's why your energy, your nerve comfort, and your everyday wellness are all connected to one molecule.

Romira R-ALA. 600mg. Pure R-form. Organic Coconut Oil.

60-day guarantee. Made in the USA."""
    ),
    (
        "6. The Age Angle",
        """Nobody tells you this about getting older:

Your body's natural antioxidant defenses start to decline. Oxidative stress builds up faster. Nerves — which depend on huge amounts of energy and protection — become more vulnerable.

That's when the tingling starts. The burning at night. The sensitivity that wasn't there ten years ago.

R-Alpha Lipoic Acid is one of the most studied nutrients for healthy aging because it hits multiple systems at once: antioxidant defense, mitochondrial energy, glucose metabolism, nerve comfort.

It even regenerates Vitamin C, Vitamin E, and Glutathione — amplifying your entire antioxidant network from a single daily supplement.

Romira R-ALA. Pure R-form. 600mg. Organic Coconut Oil. 60-day guarantee."""
    ),
    (
        "7. Pattern Interrupt",
        """Stop scrolling.

If you have nerve discomfort — burning, tingling, pins and needles, sensitivity in your feet or hands — this is the most important thing you'll read today.

Most people try magnesium. Some try CBD. Some try turmeric. Very few try the one nutrient that actually targets the root of nerve oxidative stress.

R-Alpha Lipoic Acid. The universal antioxidant. Water and fat soluble. Works inside and outside your cells. Supports mitochondrial energy. Supports healthy blood sugar. Regenerates other antioxidants.

And it comes in two forms. Most brands use the cheap mixed version. Romira uses the pure R-form — the one your body is built to work with.

600mg. Organic Coconut Oil for absorption. 60-day guarantee."""
    ),
    (
        "8. The Blood Sugar Connection",
        """Most people don't know this: blood sugar balance and nerve health are deeply connected.

When glucose metabolism is off — even slightly — nerves are one of the first things to feel it. Tingling. Burning. Sensitivity. Discomfort that gets worse at night.

R-Alpha Lipoic Acid supports the body's natural ability to use glucose efficiently. It also fights oxidative stress directly inside nerve cells — the kind of protection that vitamin C and magnesium simply can't provide.

It's why ALA has been used in clinical settings in Europe for decades as a nerve-support nutrient.

Romira R-ALA. 600mg. Pure R-form. Organic Coconut Oil. Made in the USA. 60-day guarantee."""
    ),
    (
        "9. The Skeptic",
        """You've probably tried supplements before. And been disappointed.

Fair.

Here's what's different about R-Alpha Lipoic Acid: it doesn't just mask discomfort. It supports the actual systems connected to nerve health — antioxidant protection, mitochondrial energy, blood sugar metabolism, circulation.

And here's what's different about the R-form specifically: it's the form your body produces naturally. Not the synthetic S-form mixed in with most ALA supplements. Just the molecule your body was built to use.

Romira. Pure R-form. 600mg. Organic Coconut Oil for absorption. 60-day money-back guarantee — because we know you've been burned before."""
    ),
    (
        "10. The Comparison",
        """Vitamin C: water-soluble only.
Vitamin E: fat-soluble only.
Magnesium: doesn't touch nerve oxidative stress.
CBD: unregulated, inconsistent quality.
Turmeric: poor bioavailability without special formulation.

R-Alpha Lipoic Acid: water AND fat soluble, works inside and outside cells, supports mitochondrial energy, supports blood sugar, regenerates Vitamin C + E + Glutathione.

One molecule. Multiple systems. For people who want to support their nerves properly — not just throw darts.

Romira R-ALA. 600mg. Pure R-form. Organic Coconut Oil. 60-day guarantee."""
    ),
]

for title_text, body_text in ads:
    doc.add_heading(title_text, level=2)
    p = doc.add_paragraph(body_text)
    p.paragraph_format.space_after = Pt(6)
    doc.add_paragraph('')

# Top picks note
doc.add_heading('TOP PICKS', level=1)
recs = [
    ("Headlines to test first:", "#1 and #10 — highest scroll-stop, speak directly to symptoms and past failures."),
    ("Primary texts to test first:", "#1 (The Night Story) — cold audiences. #4 (Short & Punchy) — mobile/video. #9 (The Skeptic) — retargeting."),
]
for label, note in recs:
    p = doc.add_paragraph()
    p.add_run(label + " ").bold = True
    p.add_run(note)

output_path = r"C:\Users\Darin Game\Downloads\Romira_R-ALA_Facebook_Ads.docx"
doc.save(output_path)
print(f"Saved to: {output_path}")
