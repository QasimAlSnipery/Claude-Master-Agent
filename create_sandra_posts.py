from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading('Sandra Hughes — Facebook Posts (Full Version)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

posts = [
    {
        'number': 1,
        'topic': 'Morning Back Pain',
        'headline': 'My back finally stopped ruining my mornings',
        'image': 'Candid UGC-style photo of a British woman in her mid-50s sitting on the edge of an unmade bed, early morning grey window light coming from one side. She is wearing a faded pink cotton nightshirt, hair slightly dishevelled from sleep. One hand pressed firmly against her lower back, the other gripping the mattress edge. She is looking down at the floor with a tight-lipped grimace — mouth closed, brow furrowed. On the nightstand: a rectangular weekly pill organiser, a glass of water with a straw, and a tube of Voltaren gel. Mismatched floral bedding slightly tangled. Warm yellow lamp on, grey daylight from window creating uneven lighting. Slight motion blur. Tilted horizon. Looks like she took it herself with her phone propped against the lamp. Photorealistic, authentic, no studio lighting. 1:1 square.',
        'primary': 'For 4 years I dreaded waking up. The first 10 minutes out of bed were agony — my lower back just locked solid every single night.\n\nPhysio, heat pads, ibuprofen, that expensive foam mattress my daughter talked me into. Nothing worked for more than a day or two.\n\nA lady in my walking group mentioned a natural supplement she had been taking. I was sceptical honestly. I had heard it all before.\n\nThree weeks in I noticed I was getting up without grabbing the bedpost. By week five I was downstairs making tea before my husband even woke up.\n\nI am not saying it is a miracle. I am saying it is the first thing in four years that actually made a difference to my mornings.\n\nI found an article that explains exactly what it is and how it works — leaving the link below if you want to have a read and see if it might help you too. 👇'
    },
    {
        'number': 2,
        'topic': 'Knee Pain / Stairs',
        'headline': 'I can do the stairs again. Here\'s why.',
        'image': 'UGC-style photograph of a British woman aged 54-58 gripping a wooden stair banister with both hands, mid-way up a carpeted staircase in a modest English home. She is wearing navy blue capri leggings and an oversized grey cardigan. Her face is turned slightly downward, eyes closed in a tight wince, mouth compressed. The staircase carpet is slightly worn, cream coloured. Natural window light from the landing above. A pair of slip-on shoes visible at the bottom of the stairs. The image is cropped from the side, slightly off-centre, as if caught by a family member on a phone. Warm domestic lighting, slight grain. Photorealistic, no studio lighting. 1:1 square.',
        'primary': 'Six months ago I was genuinely considering moving my bedroom downstairs. One hand on the banister, one step at a time, trying not to let my husband see how bad it had got.\n\nBone on bone, the doctor said. Basically sent me home with a leaflet.\n\nI tried glucosamine, turmeric, the copper bracelet my mum swore by. I did the physio exercises every morning for three months straight. Still the same grinding ache every time I bent my knees.\n\nThen my sister — retired nurse — told me about a supplement she had started recommending to people. Said the research behind it was solid.\n\nI gave it 30 days. By day 25 I walked up the stairs carrying a laundry basket and did not even think about it until I got to the top.\n\nThere is a full article explaining what it actually does — link is in the comments if you want to take a look. 👇'
    },
    {
        'number': 3,
        'topic': 'Painsomnia / 3AM',
        'headline': 'I slept through the night for the first time in 2 years',
        'image': 'Dimly lit bedroom scene at 3am. A British woman in her mid-50s sitting upright in bed against the headboard, knees pulled slightly up. She is wearing an oversized faded white sleep shirt. The only light source is a warm yellow bedside lamp creating soft uneven shadows. Her face shows exhausted resignation — eyes half-open and heavy, slight frown, mouth loosely closed. One hand resting on her knee. On the nightstand: a digital clock showing 3:14am, a glass of water, and a small amber pill bottle. Bedsheets tangled and pushed to one side. The room is slightly cluttered — reading glasses on the nightstand, a paperback face-down. Photorealistic, intimate, slightly grainy as if taken on an older phone. 1:1 square.',
        'primary': 'I used to dread bedtime. The aching in my back and the burning in my feet would start up the minute I lay still and that was it — awake at 2, awake at 3, awake at 4.\n\nI was exhausted all day and in pain all night. My doctor told me to try to rest more. I nearly laughed in his face.\n\nI came across an article about a natural supplement that apparently helps with the kind of nerve and joint pain that gets worse at night. Thought — well, I have tried everything else.\n\nThe first week nothing dramatic. But by week three I was waking up and realising it was morning. Actually morning. I had slept.\n\nI cried. Genuinely. Because if you have had painsomnia you know what it means to just sleep.\n\nLeaving the article in the comments — it explains it far better than I can. Worth a read if the nights are hard for you too. 👇🙏'
    },
    {
        'number': 4,
        'topic': 'Doctor Frustration',
        'headline': 'My doctor had no answers. This did.',
        'image': 'UGC-style photo of a British woman in her mid-50s sitting at a plain wooden kitchen table, a cup of tea in front of her going cold. She is wearing a muted olive green jumper, hair pulled back loosely. Her elbow is on the table, forehead resting in one hand, eyes downcast with a tired, defeated expression. On the table beside the tea: a printed NHS appointment letter, a small notebook with handwritten notes. Modest English kitchen background — cream cabinets, a tea towel hanging on the oven handle, a calendar on the wall. Soft overcast window light. Slight lens distortion at edges. The image looks like it was taken on a phone propped against the fruit bowl. Photorealistic. 1:1 square.',
        'primary': 'Came back from the GP this morning for the third time this year. Joint pain, burning legs at night, barely sleeping.\n\nSandra, you are 56. This is normal at your age.\n\nNormal. Right.\n\nI know my own body. I knew something was not right and I was not prepared to just accept it.\n\nI did my own research. Found an article about a natural compound — not a painkiller, not a prescription — that apparently works on the root cause of the kind of inflammation and nerve pain I was describing.\n\nI was sceptical. But I was also desperate.\n\nTwo months on — I am sleeping better, moving better, and I have not needed to take ibuprofen every single day for the first time in three years.\n\nI am not a doctor and I cannot promise it will work for you. But it worked for me when nothing else did.\n\nThe article that started it all is linked below — have a read and make your own mind up. 👇'
    },
    {
        'number': 5,
        'topic': 'Gardening Loss',
        'headline': 'I got my garden back.',
        'image': 'Quiet English garden scene, late afternoon soft light. In the foreground, an empty wooden garden kneeler pad abandoned on the grass next to an unfinished flower bed with a trowel resting in the soil. In the soft background, slightly out of focus, a British woman in her mid-50s stands at the back door of a modest brick terraced house, one hand on the door frame, looking out at the garden. She is wearing a faded floral apron over a cotton blouse, gardening gloves still on her hands. Her posture suggests she had to stop — one hand pressing against her lower back. The garden looks slightly overgrown at the edges. Overcast English sky, warm muted tones. The whole image feels like a quiet, private moment of loss. Phone-snapped quality, slightly blurry. Photorealistic. 1:1 square.',
        'primary': 'Last summer I had to leave the trowel in the soil and go sit down. Back and knees just gave out after 20 minutes. I had a little cry about it if I am honest.\n\nGardening was mine. It was my Saturday, my peace, my thing. And the pain was slowly taking it away.\n\nI tried everything. Physio, anti-inflammatories, the lot. Some things helped a little. Nothing gave me my garden back.\n\nA friend sent me an article about a supplement she had been reading about. Natural, not a painkiller, works differently to the usual things. I read it three times before I ordered it.\n\nThis past Saturday I spent two hours out there. Weeded the whole front bed. Planted the geraniums. Made it back inside under my own steam, no drama.\n\nIt was not like that overnight — took about five weeks. But something has genuinely shifted.\n\nLeaving the article in the comments — it is worth a read if pain has taken something you love from you too. 👇❤️'
    },
    {
        'number': 6,
        'topic': 'Grandkids',
        'headline': 'I sat on the floor with my granddaughter. I cried after.',
        'image': 'Warm living room scene in a modest English home. A British woman in her mid-50s is seated in a faded floral armchair, slightly leaning forward with both hands resting on her knees. In the foreground on the carpet, just out of easy reach, a childs colourful puzzle is partially assembled. The womans expression is soft but tinged with quiet frustration — she wants to get down there but cannot. She is wearing a soft lilac cardigan over a white blouse. The room has warm afternoon light — family photos on the wall behind her, a small bookshelf, typical English domestic interior. The image is taken from a mid-angle as if someone else is in the room with her. Warm, intimate, slightly imperfect. Photorealistic. 1:1 square.',
        'primary': 'Evie is 4. For the past year I have been watching her play from the armchair because getting down on the floor and back up again was just too much.\n\nI hated it. I hated that she was starting to just accept that Nana could not do certain things.\n\nI had tried everything my GP suggested. Tried the supplements from Holland and Barrett. Tried the exercises. Tried resting more. Nothing moved the needle enough.\n\nThen I read an article about a natural compound that a lot of people with joint stiffness had been quietly using. It works differently to the usual anti-inflammatories — acts on the pain at a different level apparently.\n\nFive weeks after starting it, Evie asked me to sit with her and do her puzzle. And I got down on that floor, did the puzzle, and got back up.\n\nI went to the kitchen and had a little cry. A happy one.\n\nLink to the article is below — if getting on the floor with your grandkids matters to you, it is worth five minutes of your time. 👇🙏'
    },
    {
        'number': 7,
        'topic': 'Burning Feet / Neuropathy',
        'headline': 'The burning in my feet finally stopped waking me up',
        'image': 'Close-up bedroom scene at night. A British woman in her mid-50s sitting on the edge of the bed, both feet flat on the carpet. She is leaning forward, both hands wrapped around one foot, rubbing it gently. She is wearing a loose nightshirt and bed socks that have been pushed down around her ankles. The bedside lamp is on, casting warm yellow light. Her face is partially visible — brow furrowed, eyes downcast, expression of exhausted discomfort. On the nightstand: a tube of cream, a glass of water, a phone screen showing 2:47am. The image is cropped mid-body, the face not fully showing, making it feel universal. Slightly grainy, warm tones, intimate and real. Photorealistic. 1:1 square.',
        'primary': 'Like someone holding a lighter to the bottom of my feet. Every single night from about midnight onwards.\n\nI had tried the creams, the cool water, elevating them, keeping them out of the duvet. Nothing stopped it. My doctor said it was nerve related and that was basically that.\n\nI was in a Facebook group for people with nerve pain and someone mentioned a natural supplement they had been taking — said it had made a real difference to the nighttime burning specifically.\n\nI was not hopeful. But I was desperate.\n\nBy week four I was sleeping through to 5am without waking up in burning pain. That might not sound like much but after two years of 2am agony it felt enormous.\n\nIt is not gone completely. But it is manageable now in a way it was not before.\n\nI found a proper article that explains what it is and the science behind it — leaving it in the comments below. If the nights are unbearable for you, please have a read. 👇'
    },
    {
        'number': 8,
        'topic': 'Tried Everything',
        'headline': 'After £600 of trying everything — this finally worked',
        'image': 'Overhead flat-lay style shot of a British womans cluttered bedside table, taken on a phone from above. On the surface: a weekly pill organiser half-empty, a tube of Deep Heat cream, a copper magnetic bracelet, a bottle of turmeric capsules, an ice pack in a faded blue cover, a printed sheet of NHS physio exercises, and a half-finished mug of chamomile tea. One weathered hand with visible age spots and no nail varnish is reaching into frame picking up the turmeric bottle. Warm imperfect light from a bedside lamp. The table surface is slightly dusty. The whole image reads as someone who has tried everything. Photorealistic, phone-quality, no studio lighting. 1:1 square.',
        'primary': 'Turmeric. Glucosamine. Deep Heat. Physio. The copper bracelet my mum swore by. Ice. Heat. Ibuprofen. Cutting out gluten. Cutting out dairy.\n\nI kept a little notebook of everything I tried. 14 things over two years. Some helped a tiny bit. Most did nothing.\n\nI was about ready to just accept this was my life now.\n\nThen I came across an article about a natural compound I had never heard of before. Not sold in Boots, not something my GP ever mentioned. The article explained why it works differently to the usual things and why that matters for joint and nerve pain specifically.\n\nI gave it six weeks.\n\nIt is now the first thing I have ever taken off my list not because it stopped working — but because it actually worked.\n\nNot saying it will be the one for everyone. But if you have been through the same exhausting cycle, it might be worth reading about.\n\nArticle is in the comments. 👇'
    },
    {
        'number': 9,
        'topic': 'Small Win / Micro Victory',
        'headline': '62 and moving better than I did at 58',
        'image': 'Outdoor photograph of a British woman in her mid-50s walking along a quiet residential English street, late afternoon autumn light. She is wearing a mustard yellow quilted gilet over a navy long-sleeve top, comfortable walking trainers. She is mid-stride, looking slightly to the side with a small private smile — not a beaming grin, just the quiet satisfaction of someone who made it further than they expected. Slightly blurry, as if taken by someone walking behind her. Fallen leaves on the pavement. The street is ordinary — semi-detached brick houses, a parked car, a low garden wall. Warm golden light. The image feels like a real moment captured, not posed. Photorealistic. 1:1 square.',
        'primary': 'Eighteen months ago I was taking the stairs one at a time, could not walk to the end of the road without stopping, and was starting every morning in pain.\n\nI had resigned myself to it if I am honest. Thought this was just what getting older felt like.\n\nMy neighbour — she is 68 — mentioned she had been taking a supplement for about three months and had started back at her aqua aerobics class. I thought she was winding me up.\n\nShe was not. She sent me the article she had read about it and I spent an evening going through it properly.\n\nI am now on month four. I walked to the park and back yesterday — just under two miles. I have not done that in two years.\n\nI am not back to where I was at 40 and I am not pretending otherwise. But the difference in how I move and how I feel day to day is real.\n\nLinking the article below — it explains everything far better than I can. 👇❤️'
    },
    {
        'number': 10,
        'topic': 'Weather & Joint Pain',
        'headline': 'Rainy days don\'t beat me up anymore',
        'image': 'Interior window scene in a modest English living room. Rain streaking down the window glass, grey overcast sky outside. A British woman in her mid-50s is seated sideways on a sofa next to the window, legs tucked up slightly, both hands wrapped around a large mug of tea. She is wearing a thick oatmeal-coloured knit jumper and thick socks. Her gaze is out the window, expression quietly resigned — not dramatic, just tired. One hand occasionally rubs her knee. On the sofa beside her: a folded blanket, a TV remote. The room is warm and dim — a standard lamp on, rain-light from the window. Net curtains partially drawn. Very English, very domestic, very real. Photorealistic, warm tones, phone quality. 1:1 square.',
        'primary': 'My knees have always known when rain was coming. Better than any forecast. The deep aching would start the afternoon before and by the time the clouds broke I would be useless.\n\nCold fronts, damp weather, anything grey — my joints felt every bit of it. I just accepted it as part of life after 50.\n\nSix months ago I started taking a natural supplement after reading an article about it. I was not expecting much — I had tried plenty before.\n\nLast week we had three days of solid rain here. The knees knew it was coming — there was that familiar warning twinge on the Tuesday.\n\nBut it did not land the way it used to. The full ache never arrived. I made it through the week without a single day on the sofa.\n\nSmall thing maybe. But if you have spent years being flattened by the weather you will understand why I am telling everyone I know.\n\nArticle explaining what I take and how it works is in the comments below. 👇☕'
    }
]

for post in posts:
    doc.add_page_break()

    doc.add_heading(f"POST {post['number']} — {post['topic']}", level=1)

    doc.add_heading('HEADLINE', level=2)
    p = doc.add_paragraph(post['headline'])
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(13)

    doc.add_paragraph()

    doc.add_heading('IMAGE PROMPT', level=2)
    doc.add_paragraph(post['image'])

    doc.add_paragraph()

    doc.add_heading('PRIMARY TEXT (Post Caption)', level=2)
    doc.add_paragraph(post['primary'])

    doc.add_paragraph()

output_path = r'C:\Users\moham\OneDrive\Desktop\Sandra_Hughes_Facebook_Posts.docx'
doc.save(output_path)
print('Done. File saved to:', output_path)
