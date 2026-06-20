# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt

doc = Document()

def H1(t): return doc.add_heading(t, level=1)
def H2(t): return doc.add_heading(t, level=2)
def H3(t): return doc.add_heading(t, level=3)
def para(t, bold=False, italic=False, size=11):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.bold = bold; r.italic = italic; r.font.size = Pt(size); return p
def bullet(t):
    p = doc.add_paragraph(style='List Bullet'); p.add_run(t); return p
def kv(label, val):
    p = doc.add_paragraph(); r = p.add_run(label + ": "); r.bold = True; p.add_run(val); return p

doc.add_heading('Alpha-Lipoic Acid - UGC Ad Production File', level=0)
para('Master prompt-generation brief for ChatGPT -> 10-second AI video prompts', italic=True, size=12)
para('3 scripts | 3 distinct middle-aged women | Vertical 9:16 UGC | Voiceover pacing: 150 words per minute', size=10)
doc.add_paragraph()

H1('READ FIRST - Instructions for ChatGPT')
para('You are a senior AI-video prompt engineer. Your job: turn each SCENE in this document into a fully detailed, copy-paste-ready 10-second video generation prompt (for tools like Sora, Veo, Kling, or Runway). Produce ONE prompt per scene, in order, grouped by script.', size=11)
doc.add_paragraph()
para('Every 10-second prompt you generate MUST include all of the following, written out explicitly (never assume the model remembers a previous scene - repeat the character description in EVERY prompt):', bold=True)
bullet('CHARACTER: full physical appearance of the woman (age, build, hair, skin, face, hands/feet if shown) - copied from her Character Sheet below, every time.')
bullet('WARDROBE: exact clothing, jewelry, glasses for that woman.')
bullet('SETTING: the room/location, props, and time of day.')
bullet('LIGHTING: natural soft / window light / lamp, etc. - keep UGC-realistic, not studio-glossy.')
bullet('CAMERA: framing (selfie arm-length, chest-up, etc.), angle, lens feel, subtle handheld movement. Phone-shot look, NOT cinematic polish.')
bullet('ACTION and EXPRESSION: what she physically does and her facial emotion during the line (rubbing her foot, leaning in, sighing, light laugh).')
bullet('B-ROLL / INSERTS: any cutaway the scene calls for (close-up of hands tingling, product bottle, link tap) described visually.')
bullet('ON-SCREEN TEXT (captions): the burned-in caption text for that scene, bold, high-contrast, bottom-center, sound-off readable.')
bullet('SPOKEN VOICEOVER: the exact VO line(s) for that scene, in quotes. Target ~25 words per 10s so delivery lands at 150 words per minute.')
doc.add_paragraph()
para('Global rules to bake into every prompt:', bold=True)
bullet('Aspect ratio: 9:16 vertical. Mobile-first.')
bullet('Style: authentic UGC / talking-head selfie. Slightly imperfect, real, relatable. NOT a commercial, NOT studio lighting, NOT a model. She must look like a real 50-something woman filming herself.')
bullet('Pacing: 150 words per minute = ~25 spoken words per 10-second clip. Natural, conversational delivery.')
bullet('Each woman must look clearly DIFFERENT from the other two (different ethnicity, hair, build, vibe, home).')
bullet('Compliance: keep claims safe - "supports nerve health", "oxidative stress", "helps support healthy nerve function". NEVER claim it cures or treats disease.')
bullet('CTA scenes must visually emphasize the 50% OFF / limited / ending-soon offer with an on-screen text badge.')
doc.add_paragraph()
para('Output format you should return: For each scene, label it (e.g. "SCRIPT 1 - SCENE 3"), then give the finished prompt as one detailed paragraph, plus a separate "On-screen text:" line and "Voiceover:" line.', italic=True)
doc.add_page_break()

H1('Character Sheets - the 3 women')

H2('Woman A - "Linda", 54 - warm and tired-but-hopeful')
kv('Vibe', 'Soft-spoken, kind, talks like she is telling a close friend over coffee. Gentle, sincere, a little weary.')
kv('Ethnicity / skin', 'White, fair skin with natural age lines, light freckles, minimal makeup.')
kv('Build', 'Average build, soft features, warm genuine smile.')
kv('Hair', 'Shoulder-length silver-blonde, loose natural waves, slightly tucked behind one ear.')
kv('Face', 'Soft round face, light crow’s-feet, reading glasses pushed up on her head.')
kv('Wardrobe', 'Cozy cream knit cardigan over a soft pastel top, small stud earrings, simple wedding band.')
kv('Setting', 'Bright home kitchen or breakfast nook, morning light, mug of tea on the counter, plants in background.')
kv('Camera feel', 'Arm-length selfie, chest-up, very subtle handheld sway, warm natural window light.')

H2('Woman B - "Carol", 58 - blunt, no-nonsense skeptic')
kv('Vibe', 'Dry, direct, arms-crossed energy, "I do not believe hype." Wins trust by being unimpressed.')
kv('Ethnicity / skin', 'White, cooler skin tone, visible age lines, essentially no makeup.')
kv('Build', 'Lean, sharper features, strong jaw, level steady gaze.')
kv('Hair', 'Short salt-and-pepper pixie / cropped cut, neat.')
kv('Face', 'Angular face, thin-rim reading glasses worn on the nose, slight raised-eyebrow skepticism.')
kv('Wardrobe', 'Crisp navy blouse or button-down, thin silver necklace, watch.')
kv('Setting', 'Tidy living room, neutral tones, sitting in a fabric armchair, bookshelf behind her, soft daylight.')
kv('Camera feel', 'Phone propped on a stand, chest-up, mostly static with tiny natural movement, even daylight.')

H2('Woman C - "Maria", 51 - energetic and expressive')
kv('Vibe', 'Animated, passionate, lots of hand gestures, "let me TELL you." Fast, warm, expressive.')
kv('Ethnicity / skin', 'Latina, warm golden-brown skin, healthy glow, natural everyday makeup.')
kv('Build', 'Curvier build, lively, expressive eyebrows and hands.')
kv('Hair', 'Dark wavy shoulder-length hair with subtle grey streaks, volume, moves as she talks.')
kv('Face', 'Oval expressive face, bright eyes, gold hoop earrings.')
kv('Wardrobe', 'Colorful blouse or a draped scarf (coral / teal), layered necklace.')
kv('Setting', 'Cozy bright living room, warm lamps, throw pillows, family photos softly blurred behind.')
kv('Camera feel', 'Arm-length selfie, more movement, she leans in/out, gestures into frame, warm indoor light.')
doc.add_page_break()

def script_block(num, persona, hook_note, scenes):
    H1('SCRIPT %d - %s' % (num, persona))
    para('HOOK IS LOCKED - do not change the opening line. %s' % hook_note, bold=True)
    para('Voiceover pacing: 150 wpm (~25 words per 10-second scene). Total %d scenes.' % len(scenes), italic=True, size=10)
    doc.add_paragraph()
    for i, sc in enumerate(scenes, 1):
        H3('Scene %d  (0:%02d-0:%02d)' % (i, (i-1)*10, i*10))
        kv('Voiceover (%d words)' % len(sc['vo'].split()), '"%s"' % sc['vo'])
        kv('On-screen text', sc['text'])
        kv('What we see (action and shot)', sc['see'])
        kv('B-roll / insert', sc['broll'])
        doc.add_paragraph()
    doc.add_page_break()

s1 = [
 {'vo':"If you have burning, tingling, or numbness in your hands or feet, pay attention. For two years, I woke up at 3am with my feet on fire.",
  'text':"Burning. Tingling. Numbness?",
  'see':"Linda close, selfie, looking straight into camera, serious and sincere. Slight lean in on 'pay attention.' Then a soft pained expression remembering the nights.",
  'broll':"Quick 1-sec cut: her hand gently rubbing the top of her foot at the edge of a bed, dim night lighting."},
 {'vo':"I tried everything just to get through the day. But nothing touched it at night. Here is what nobody told me about nerve discomfort.",
  'text':"Nobody told me this...",
  'see':"Linda in kitchen, slight sigh, shakes head, then eyebrows lift on 'here is what nobody told me' - curiosity building.",
  'broll':"None - stay on her face for connection."},
 {'vo':"It isn't only about pain. It's tied to oxidative stress and how well your nerves are actually supported. That's when I found alpha-lipoic acid.",
  'text':"It's about NERVE SUPPORT",
  'see':"Linda explaining calmly, gentle hand gesture. Small nod of realization on 'that is when I found.'",
  'broll':"Soft insert: a simple supplement bottle on the kitchen counter, morning light."},
 {'vo':"An antioxidant studied for nerve health that helps support healthy nerve function. I'm not saying it's a miracle, and don't stop anything your doctor gave you.",
  'text':"Alpha-Lipoic Acid - studied for nerve health",
  'see':"Linda holds bottle lightly, honest reassuring tone, open palm on 'not a miracle' - builds trust.",
  'broll':"Close-up of the bottle label held in her hands."},
 {'vo':"But if your hands or feet burn or go numb, you need to see this. It's 50% off right now, ending soon. Tap the link before it's gone.",
  'text':"50% OFF - ENDING SOON. Tap the link",
  'see':"Linda warm encouraging smile, points down toward link, urgency but kind. Direct eye contact on CTA.",
  'broll':"Insert: finger tapping a phone screen / link, plus animated 50% OFF badge."},
]
script_block(1, '"Linda", 54 - warm and hopeful', 'Opening line: "If you have burning, tingling, or numbness in your hands or feet, pay attention."', s1)

s2 = [
 {'vo':"Most people with nerve discomfort are frustrated because they feel like they're only managing the symptoms. I was that person.",
  'text':"Just managing symptoms?",
  'see':"Carol static in armchair, level skeptical gaze, slight raised eyebrow, dry delivery. Points at herself on 'I was that person.'",
  'broll':"None - hold on her unimpressed expression for authenticity."},
 {'vo':"Take something, wait, feel a little better. But the burning comes back, the numbness is still there. I was sick of just covering it up.",
  'text':"It keeps coming back.",
  'see':"Carol counts on fingers ('take something, wait'), then flat annoyed look on 'comes back.' Genuine frustration.",
  'broll':"Quick insert: hand flexing fingers, subtle tingle visual."},
 {'vo':"Here's what changed my mind. Your nerves are sensitive, and oxidative stress affects how they feel and function. So I looked at what researchers point to.",
  'text':"What changed my mind",
  'see':"Carol leans forward slightly, more engaged now, matter-of-fact teaching tone. Pushes glasses up.",
  'broll':"None."},
 {'vo':"Not some complicated drug. Alpha-lipoic acid. An antioxidant that supports your body's defense against oxidative stress. This isn't medical advice, so talk to your doctor.",
  'text':"Alpha-Lipoic Acid",
  'see':"Carol holds bottle, gives it a no-nonsense glance and nod of approval. Honest shrug on 'talk to your doctor.'",
  'broll':"Close-up of bottle held in steady hand."},
 {'vo':"But if you're tired of masking it, look into this. It's 50% off and ending soon. The link's below. Don't wait on this one.",
  'text':"50% OFF - LIMITED. Link below",
  'see':"Carol points firmly down at link, direct deadpan urgency, single nod. Trust through bluntness.",
  'broll':"Insert: 50% OFF badge + finger tapping link."},
]
script_block(2, '"Carol", 58 - blunt skeptic', 'Opening line: "Most people with nerve discomfort are frustrated because they feel like they are only managing the symptoms."', s2)

s3 = [
 {'vo':"The reason your hands or feet may be tingling, burning, or going numb is not always as simple as you're getting older.",
  'text':"It's NOT just 'getting older'",
  'see':"Maria animated selfie, expressive eyebrows, hand gesture air-quoting 'getting older.' High energy hook.",
  'broll':"None - stay on her expressive face."},
 {'vo':"I kept telling myself, maybe I slept wrong, maybe it's my shoes, maybe it'll go away. It did not go away.",
  'text':"'Maybe it'll go away...' (it didn't)",
  'see':"Maria lists excuses on fingers with shrugs and head tilts, then firm shake of head on 'did not go away.'",
  'broll':"Quick insert: rubbing her hand/wrist, mild discomfort look."},
 {'vo':"It started happening at night, sitting down, when I tried to relax. So here's the part no one explains.",
  'text':"The part no one explains",
  'see':"Maria leans in close to camera, lowers voice slightly, conspiratorial 'let me tell you' energy.",
  'broll':"Soft insert: evening living room, her shifting on couch trying to get comfortable."},
 {'vo':"Researchers keep tying nerve discomfort to oxidative stress. And one antioxidant has been talked about for years for nerve health. Alpha-lipoic acid.",
  'text':"Alpha-Lipoic Acid",
  'see':"Maria bright 'aha' expression, points up on 'one antioxidant,' presents bottle proudly.",
  'broll':"Insert: bottle revealed in her hands, warm lamp light."},
 {'vo':"It helps support healthy nerve function and fights oxidative stress. It's not a cure, not a replacement for your doctor.",
  'text':"Supports healthy nerve function",
  'see':"Maria sincere for a beat, honest open-palm gesture on 'not a cure' - keeps it credible.",
  'broll':"Close-up of label."},
 {'vo':"But if your hands or feet are buzzing and burning, you should know about this. It's 50% off, limited, ending soon. Tap the link now.",
  'text':"50% OFF - LIMITED, ENDING SOON. Tap now",
  'see':"Maria big encouraging energy, points down enthusiastically, urgent excited smile on CTA.",
  'broll':"Insert: animated 50% OFF badge + finger tapping link."},
]
script_block(3, '"Maria", 51 - energetic and expressive', 'Opening line: "The reason your hands or feet may be tingling, burning, or going numb is not always as simple as you are getting older."', s3)

H1('Appendix - Quick reference for ChatGPT')
bullet('All 3 scripts share the same structure: locked hook -> personal pain -> reframe (oxidative stress) -> reveal alpha-lipoic acid -> safe disclaimer -> 50% off CTA.')
bullet('Each scene = one 10-second vertical clip. Generate prompts scene-by-scene in order.')
bullet('Always restate the woman full appearance + wardrobe + setting in every individual prompt.')
bullet('Keep it UGC-real: phone-shot, natural light, imperfect, relatable - never glossy commercial.')
bullet('Captions burned in, bold, high-contrast, sound-off readable.')
bullet('CTA scenes: emphasize 50% OFF + limited/ending-soon badge on screen.')

doc.save('Alpha_Lipoic_Acid_UGC_Prompt_File.docx')
print("SAVED: Alpha_Lipoic_Acid_UGC_Prompt_File.docx")
