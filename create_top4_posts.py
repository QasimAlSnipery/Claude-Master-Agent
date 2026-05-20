from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading('Sandra Hughes — Top 4 Facebook Posts', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph('Run these 4 together. Lead with Post 6 (Grandkids).')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

posts = [
    {
        'number': '6 — LEAD POST',
        'topic': 'Grandkids (Run This First)',
        'headline': 'I sat on the floor with my granddaughter. I cried after.',
        'image': 'Warm living room scene in a modest English home. A British woman in her mid-50s is seated in a faded floral armchair, slightly leaning forward with both hands resting on her knees. In the foreground on the carpet, just out of easy reach, a childs colourful puzzle is partially assembled. The womans expression is soft but tinged with quiet frustration — she wants to get down there but cannot. She is wearing a soft lilac cardigan over a white blouse. The room has warm afternoon light — family photos on the wall behind her, a small bookshelf, typical English domestic interior. The image is taken from a mid-angle as if someone else is in the room with her. Warm, intimate, slightly imperfect. Photorealistic. 1:1 square.',
        'primary': 'Evie is 4. For the past year I have been watching her play from the armchair because getting down on the floor and back up again was just too much.\n\nI hated it. I hated that she was starting to just accept that Nana could not do certain things.\n\nI had tried everything my GP suggested. Tried the supplements from Holland and Barrett. Tried the exercises. Tried resting more. Nothing moved the needle enough.\n\nThen I read an article about a natural compound that a lot of people with joint stiffness had been quietly using. It works differently to the usual anti-inflammatories — acts on the pain at a different level apparently.\n\nFive weeks after starting it, Evie asked me to sit with her and do her puzzle. And I got down on that floor, did the puzzle, and got back up.\n\nI went to the kitchen and had a little cry. A happy one.\n\nLink to the article is below — if getting on the floor with your grandkids matters to you, it is worth five minutes of your time. 👇🙏',
        'why': 'WHY THIS IS THE LEAD: Grandkids is the single most powerful emotional anchor for this audience. Hits the deepest fear — becoming useless to the people they love most. Highest share and comment potential.'
    },
    {
        'number': '1',
        'topic': 'Morning Back Pain (Widest Reach)',
        'headline': 'My back finally stopped ruining my mornings',
        'image': 'Candid UGC-style photo of a British woman in her mid-50s sitting on the edge of an unmade bed, early morning grey window light coming from one side. She is wearing a faded pink cotton nightshirt, hair slightly dishevelled from sleep. One hand pressed firmly against her lower back, the other gripping the mattress edge. She is looking down at the floor with a tight-lipped grimace — mouth closed, brow furrowed. On the nightstand: a rectangular weekly pill organiser, a glass of water with a straw, and a tube of Voltaren gel. Mismatched floral bedding slightly tangled. Warm yellow lamp on, grey daylight from window creating uneven lighting. Slight motion blur. Tilted horizon. Looks like she took it herself with her phone propped against the lamp. Photorealistic, authentic, no studio lighting. 1:1 square.',
        'primary': 'For 4 years I dreaded waking up. The first 10 minutes out of bed were agony — my lower back just locked solid every single night.\n\nPhysio, heat pads, ibuprofen, that expensive foam mattress my daughter talked me into. Nothing worked for more than a day or two.\n\nA lady in my walking group mentioned a natural supplement she had been taking. I was sceptical honestly. I had heard it all before.\n\nThree weeks in I noticed I was getting up without grabbing the bedpost. By week five I was downstairs making tea before my husband even woke up.\n\nI am not saying it is a miracle. I am saying it is the first thing in four years that actually made a difference to my mornings.\n\nI found an article that explains exactly what it is and how it works — leaving the link below if you want to have a read and see if it might help you too. 👇',
        'why': 'WHY THIS IS IN THE TOP 4: Widest possible audience. Every 50+ person with pain knows the exact feeling of dreading mornings. Maximum cold traffic reach.'
    },
    {
        'number': '3',
        'topic': 'Painsomnia / 3AM (Highest Intent)',
        'headline': 'I slept through the night for the first time in 2 years',
        'image': 'Dimly lit bedroom scene at 3am. A British woman in her mid-50s sitting upright in bed against the headboard, knees pulled slightly up. She is wearing an oversized faded white sleep shirt. The only light source is a warm yellow bedside lamp creating soft uneven shadows. Her face shows exhausted resignation — eyes half-open and heavy, slight frown, mouth loosely closed. One hand resting on her knee. On the nightstand: a digital clock showing 3:14am, a glass of water, and a small amber pill bottle. Bedsheets tangled and pushed to one side. The room is slightly cluttered — reading glasses on the nightstand, a paperback face-down. Photorealistic, intimate, slightly grainy as if taken on an older phone. 1:1 square.',
        'primary': 'I used to dread bedtime. The aching in my back and the burning in my feet would start up the minute I lay still and that was it — awake at 2, awake at 3, awake at 4.\n\nI was exhausted all day and in pain all night. My doctor told me to try to rest more. I nearly laughed in his face.\n\nI came across an article about a natural supplement that apparently helps with the kind of nerve and joint pain that gets worse at night. Thought — well, I have tried everything else.\n\nThe first week nothing dramatic. But by week three I was waking up and realising it was morning. Actually morning. I had slept.\n\nI cried. Genuinely. Because if you have had painsomnia you know what it means to just sleep.\n\nLeaving the article in the comments — it explains it far better than I can. Worth a read if the nights are hard for you too. 👇🙏',
        'why': 'WHY THIS IS IN THE TOP 4: People awake at 3am in pain are at peak desperation and peak willingness to click. Highest buyer intent of any post in the set.'
    },
    {
        'number': '8',
        'topic': 'Tried Everything (Most Ready to Buy)',
        'headline': 'After £600 of trying everything — this finally worked',
        'image': 'Overhead flat-lay style shot of a British womans cluttered bedside table, taken on a phone from above. On the surface: a weekly pill organiser half-empty, a tube of Deep Heat cream, a copper magnetic bracelet, a bottle of turmeric capsules, an ice pack in a faded blue cover, a printed sheet of NHS physio exercises, and a half-finished mug of chamomile tea. One weathered hand with visible age spots and no nail varnish is reaching into frame picking up the turmeric bottle. Warm imperfect light from a bedside lamp. The table surface is slightly dusty. The whole image reads as someone who has tried everything. Photorealistic, phone-quality, no studio lighting. 1:1 square.',
        'primary': 'Turmeric. Glucosamine. Deep Heat. Physio. The copper bracelet my mum swore by. Ice. Heat. Ibuprofen. Cutting out gluten. Cutting out dairy.\n\nI kept a little notebook of everything I tried. 14 things over two years. Some helped a tiny bit. Most did nothing.\n\nI was about ready to just accept this was my life now.\n\nThen I came across an article about a natural compound I had never heard of before. Not sold in Boots, not something my GP ever mentioned. The article explained why it works differently to the usual things and why that matters for joint and nerve pain specifically.\n\nI gave it six weeks.\n\nIt is now the first thing I have ever taken off my list not because it stopped working — but because it actually worked.\n\nNot saying it will be the one for everyone. But if you have been through the same exhausting cycle, it might be worth reading about.\n\nArticle is in the comments. 👇',
        'why': 'WHY THIS IS IN THE TOP 4: Speaks directly to people who have already spent money trying to fix their pain and failed. These are the most motivated buyers — they just need a reason to try one more thing.'
    }
]

for post in posts:
    doc.add_page_break()

    doc.add_heading(f"POST {post['number']} — {post['topic']}", level=1)

    # Why this post
    why_para = doc.add_paragraph(post['why'])
    why_para.runs[0].bold = True
    why_para.runs[0].font.color.rgb = None

    doc.add_paragraph()

    doc.add_heading('HEADLINE', level=2)
    p = doc.add_paragraph(post['headline'])
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(13)

    doc.add_paragraph()

    doc.add_heading('IMAGE PROMPT', level=2)
    doc.add_paragraph(post['image'])

    doc.add_paragraph()

    doc.add_heading('PRIMARY TEXT (Post Caption)', level=2)
    doc.add_paragraph(post['primary'])

    doc.add_paragraph()

output_path = r'C:\Users\moham\OneDrive\Desktop\Sandra_Hughes_TOP4_Posts.docx'
doc.save(output_path)
print('Done:', output_path)
